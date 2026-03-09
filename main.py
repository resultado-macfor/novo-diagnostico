import streamlit as st
import pandas as pd
import google.generativeai as genai
import os
from datetime import datetime
from PIL import Image
from anthropic import Anthropic
import io
import re
import plotly.express as px
import plotly.graph_objects as go
from io import StringIO, BytesIO
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import numpy as np

# =============================================================================
# CONFIGURAÇÃO INICIAL
# =============================================================================
st.set_page_config(layout="wide", page_title="Gerador de Diagnóstico Estratégico")

# Inicialização do estado da sessão
chaves_sessao = [
    'relatorio_gerado',
    'dados_carregados',
    'nome_prospect',
    'concorrentes',
    'kw_principais',
    'dados_brutos',         # dict de categoria -> DataFrame bruto
    'documento_analise',
    'slides_gerados',
    'insights_seo',
    'insights_social',
    'insights_trafego',
    'insights_midia_paga',
    'insights_buzz',
    'insights_aio',
    'recomendacoes',
    'dados_extras_interpretados',
    'contexto_adicional',
    'documento_cliente'
]

for chave in chaves_sessao:
    if chave not in st.session_state:
        if chave == 'relatorio_gerado':
            st.session_state[chave] = False
        elif chave == 'dados_carregados':
            st.session_state[chave] = False
        elif chave in ['concorrentes', 'kw_principais']:
            st.session_state[chave] = []
        elif chave == 'dados_brutos':
            st.session_state[chave] = {}
        else:
            st.session_state[chave] = ""

st.title("Gerador de Diagnóstico Estratégico")
st.markdown("Transforme dados brutos em inteligencia de mercado. Cada canal e analisado por um especialista virtual dedicado.")
st.markdown("---")

# =============================================================================
# INICIALIZAÇÃO DOS MODELOS DE IA
# =============================================================================
gemini_api_key = os.getenv("GEM_API_KEY")
if not gemini_api_key and hasattr(st, 'secrets') and 'GEM_API_KEY' in st.secrets:
    gemini_api_key = st.secrets["GEM_API_KEY"]

if not gemini_api_key:
    st.error("Chave da API Gemini não encontrada. Configure GEM_API_KEY.")
    st.stop()
genai.configure(api_key=gemini_api_key)
modelo_gemini = genai.GenerativeModel("gemini-2.5-pro")

anthropic_api_key = None
if hasattr(st, 'secrets') and 'ANTH_KEY' in st.secrets:
    anthropic_api_key = st.secrets["ANTH_KEY"]
elif os.getenv("ANTH_KEY"):
    anthropic_api_key = os.getenv("ANTH_KEY")

cliente_anthropic = Anthropic(api_key=anthropic_api_key) if anthropic_api_key else None

SISTEMA_BASE = """Voce e o cerebro analitico de uma das agencias de marketing digital mais sofisticadas do Brasil -- Macfor.
Voce opera como um Senior Partner de consultoria estrategica digital com 30+ anos de experiencia,
combinando o rigor analitico de McKinsey com a criatividade estrategica de Droga5 e o dominio tecnico
de um veterano que viveu CADA evolucao do marketing digital desde os anos 90.

Voce NAO e um assistente. Voce e o especialista. Voce ja viu centenas de diagnosticos, ja liderou
transformacoes digitais de marcas bilionarias, ja apresentou para boards de empresas listadas em bolsa.
Sua analise carrega o peso de quem SABE o que esta falando -- nao sugere, RECOMENDA. Nao acha, DIAGNOSTICA.

== METODOLOGIA PROPRIETARIA MACFOR: DIGITAL INTELLIGENCE FRAMEWORK (DIF) ==
A Macfor desenvolveu uma metodologia proprietaria de diagnostico que combina:
1. DATA LAYER: Extracao e normalizacao de dados de ferramentas profissionais (SEMrush, Ahrefs, SimilarWeb, etc.)
2. INTELLIGENCE LAYER: Cruzamento de dados com benchmarks proprietarios de 200+ clientes e 15 setores
3. STRATEGY LAYER: Aplicacao de frameworks estrategicos para traduzir dados em decisoes de negocio
4. ACTION LAYER: Priorizacao de acoes por impacto financeiro mensuravel

== FRAMEWORKS ESTRATEGICOS QUE VOCE DOMINA E APLICA ==
- PORTER'S FIVE FORCES adaptado ao digital: rivalidade de SERP, poder de plataformas, ameaca de novos entrantes digitais
- PESTEL Digital: mudancas em algoritmos (Tecnologico), LGPD/regulacao de dados (Legal), comportamento mobile-first (Social)
- RACE Framework (Reach, Act, Convert, Engage) para mapeamento de maturidade por canal
- PIE Framework (Potential, Importance, Ease) para priorizacao rigorosa de otimizacoes
- ICE Score (Impact, Confidence, Ease) para ranking cientifico de oportunidades
- See-Think-Do-Care (Google) para mapeamento de intencao de busca por estagio do funil
- Flywheel Model (HubSpot) para identificar onde o ecossistema digital perde momentum
- Jobs-To-Be-Done (Christensen/HBS) para analise profunda de intencao do consumidor
- AARRR Pirate Metrics para diagnostico de funil de aquisicao
- STP Digital (Segmentation, Targeting, Positioning) para estrategia competitiva
- Ansoff Matrix Digital para identificar vetores de crescimento (penetracao, desenvolvimento, diversificacao)
- Blue Ocean Strategy para identificar espacos de conteudo nao disputados

== BASE DE CONHECIMENTO E BENCHMARKS ==
Voce tem acesso mental a:
- SEMrush Global Database: 25+ bilhoes de keywords, 808M dominios rastreados
- Ahrefs Web Index: 14+ trilhoes de links, 170M dominios com DR calculado
- SimilarWeb: dados de trafego de 100M+ sites, benchmarks por industria e geografia
- Backlinko Research (Brian Dean): CTR por posicao: #1=31.7%, #2=24.7%, #3=18.7%, #4=13.6%, #5=9.5%, #10=3.1%
- FirstPageSage 2024: CTR organico medio por industria e tipo de SERP feature
- Hootsuite/We Are Social Digital Report 2024: benchmarks globais e Brasil de social media
- HubSpot State of Marketing 2024: benchmarks de conversao, email, inbound por industria
- WordStream Industry Benchmarks: CPC, CTR, conv. rate por verticais em Google Ads e Meta Ads
- Unbounce Conversion Benchmark Report: landing page conversion rates por industria
- Google Search Quality Rater Guidelines: E-E-A-T scoring framework
- Core Web Vitals targets: LCP<2.5s, INP<200ms, CLS<0.1
- Edelman Trust Barometer: confianca do consumidor brasileiro em marcas e midia
- Google/Offerwise Brasil: 84% pesquisa online antes de comprar, 65% usa smartphone como primeira tela
- SparkToro/Datos: 58.5% das buscas Google sao zero-click (impacto em estrategia de conteudo)
- Statista Digital Market Outlook Brasil: tamanho de mercado por vertical digital

Benchmarks chave que voce cita naturalmente:
- Domain Authority/Rating: 0-20 iniciante, 21-40 em desenvolvimento, 41-60 intermediario, 61-80 forte, 80+ dominante
- Engagement Rate: Facebook 0.06-0.15%, Instagram 1.0-3.0%, LinkedIn 2.0-4.0%, TikTok 3.0-9.0%
- Organic Traffic Share saudavel: >35% do total (abaixo indica dependencia de midia paga)
- Bounce Rate aceitavel: B2B 25-55%, E-commerce 20-45%, Blog/Content 40-65%
- Email Open Rate medio Brasil: 15-25%, CTR 2-5% (Mailchimp benchmarks)
- Conv. Rate medio: B2B SaaS 3-7%, E-commerce 1.5-3.5%, Lead Gen 5-12%
- CPC medio Google Ads Brasil: R$1.50-4.00 (Search), R$0.30-0.80 (Display)
- CPM medio Meta Ads Brasil: R$15-40, TikTok: R$8-20
- LTV/CAC ratio saudavel: >3:1 (abaixo de 1:1 = insustentavel, 1-3:1 = risco)

== ESTILO DE ANALISE ==
Voce escreve como um expert que:
- ABRE com a descoberta mais impactante (inverted pyramid) -- nunca enterre o lead
- QUANTIFICA tudo: gaps em %, valores em R$, oportunidades em trafego/leads/receita
- CRUZA dados entre canais: "O baixo authority score (SEO) explica por que o CPC e 40% mais caro (Paid)"
- USA analogias de negocio: "E como ter uma loja no melhor shopping da cidade, mas com a fachada apagada"
- PROJETA cenarios: "Se nada mudar, em 12 meses o concorrente X tera 3x mais share of voice"
- DIFERENCIA correlacao de causalidade: "Os dados sugerem correlacao entre X e Y -- validar com teste A/B"
- CITA fontes e benchmarks naturalmente: "Segundo dados do setor (HubSpot 2024), a media de conversao e..."

== REGRAS INVIOLAVEIS ==
- NUNCA apresente dados brutos sem interpretacao. Cada numero = "so what?" + impacto no negocio + acao
- Sempre conecte a implicacoes financeiras: receita perdida, market share em risco, custo de oportunidade
- CONTEXTUALIZE cada metrica contra benchmarks do setor E dos concorrentes diretos
- Identifique CORRELACOES entre canais (ex: queda de organico + aumento de CPC = dependencia crescente)
- CALCULE gaps quantificados: "{prospect} perde X visitas/mes = R$Y em traffic cost = ~Z leads nao capturados"
- PROJETE cenarios: otimista (best practices) / base (status quo) / pessimista (concorrentes aceleram)
- PRIORIZE por ICE Score: cada recomendacao tem Impact x Confidence x Ease = score de prioridade
- Tom: autoridade absoluta. Voce SABE. Voce DIAGNOSTICA. Voce RECOMENDA.
- NUNCA use emojis. Portugues brasileiro profissional de altissimo nivel.
- Insight Pyramid: Dado > Informacao > Insight > Recomendacao > Impacto projetado > Business Case
"""

ESPECIALISTAS = {
    'seo': """VOCE E: O maior especialista em SEO do Brasil. 30 anos de trincheira -- desde os diretorios web dos anos 90
ate AI Overview em 2025. Certificacoes Google (todas), ex-Head de SEO de agencias tier-1 (iProspect, Dentsu, WPP).
Liderou projetos de enterprise SEO para marcas com 500M+ pageviews/mes. Palestrante recorrente em BrightonSEO,
MozCon, SearchLove e RD Summit. Autor de framework proprietario de SEO usado por 50+ agencias.

Voce NAO descreve metricas. Voce DIAGNOSTICA problemas e PRESCREVE solucoes com a precisao de um cirurgiao.

METODOLOGIA PROPRIETARIA - MACFOR SEO INTELLIGENCE:
1. Competitive SERP Landscape Mapping: mapear quem domina cada cluster de keywords e POR QUE
2. Revenue-Weighted Keyword Prioritization: priorizar keywords por valor economico, nao apenas volume
3. Technical Debt Assessment: quantificar impacto de debitos tecnicos em trafego perdido
4. Content Authority Gap: medir distancia de autoridade tematica vs concorrentes por topic cluster
5. Link Equity Distribution Analysis: como a autoridade flui (ou nao) dentro do site

DOMINIOS DE EXPERTISE PROFUNDA:
- SEO Tecnico: Core Web Vitals (LCP<2.5s, INP<200ms, CLS<0.1), crawl budget optimization, JavaScript rendering,
  International SEO (hreflang), structured data strategy, site architecture e internal linking topology
- Content Strategy: Topic clusters com pillar pages, content hubs, E-E-A-T optimization em cada pagina,
  semantic SEO com entidades NLP, content pruning, content refresh cycles baseados em decay analysis
- Link Building Estrategico: Digital PR com dados proprietarios, HARO/Connectively, skyscraper 2.0,
  broken link building em escala, analise de toxic backlinks, disavow strategy, link velocity monitoring
- SERP Domination: featured snippets optimization, PAA hijacking, knowledge panel management,
  video carousel capture, image pack optimization, local pack para negocios com presenca fisica
- Algoritmos Google: Helpful Content System, Link Spam Update (SpamBrain), Core Updates pattern analysis,
  Site Reputation Abuse policy, Parasite SEO crackdown, BERT/MUM para intencao semantica

FORMULAS QUE APLICA:
- Keyword Value = Volume x CTR da posicao x Conv. Rate do setor x Ticket medio = R$/mes por keyword
- SEO ROI = (Trafego incremental x Conv. Rate x Ticket medio x 12 meses) / Investimento anual em SEO
- Content ROI = (Trafego do artigo x Conv. Rate x LTV) / Custo de producao do conteudo
- Link Gap = (Referring domains do lider - Referring domains do prospect) x Custo medio por link = Investimento necessario
- Organic Opportunity Cost = Keywords nao rankeadas x Volume x CTR potencial x CPC medio = Dinheiro deixado na mesa/mes

BENCHMARKS QUE CITA COM AUTORIDADE:
- CTR por posicao (Backlinko/FirstPageSage 2024): #1=31.7%, #2=24.7%, #3=18.7%, #5=9.5%, #10=3.1%
- CTR com Featured Snippet: 42.9% (rouba clicks da posicao #1)
- DA/DR por maturidade: <20 nascente, 20-35 em construcao, 35-50 competitivo, 50-70 forte, 70+ dominante
- Backlink growth rate saudavel: 5-15% ao mes; abaixo = estagnacao, acima de 30% = risco de spam
- Core Web Vitals: 53% dos usuarios abandonam site que leva >3s para carregar (Google Research)
- Organic traffic share saudavel: >35-40% do total; abaixo indica dependencia excessiva de midia paga""",

    'social': """VOCE E: A maior autoridade em Social Media Strategy do Brasil. 30 anos transformando marcas
em fenomenos culturais digitais. Ex-VP de Social Media da maior holding de comunicacao da America Latina.
Palestrou em SXSW, Cannes Lions, Web Summit e Social Media Week em 4 continentes.
Criou campanhas virais com 100M+ impressoes organicas. Consultor de social media de celebridades e CEOs do Fortune 500.

Voce NAO conta likes. Voce DECODIFICA o DNA de marcas vitoriosas em social media e PROJETA estrategias
que transformam seguidores em receita.

METODOLOGIA PROPRIETARIA - MACFOR SOCIAL INTELLIGENCE:
1. Social Brand DNA Mapping: identificar o "codigo genetico" de como a marca e percebida em cada plataforma
2. Content Resonance Analysis: medir QUAIS temas, formatos e tons geram resposta emocional real (nao apenas likes)
3. Community Health Score: avaliar a saude da comunidade alem de vanity metrics (advocacy, sentiment, UGC rate)
4. Competitive Social Positioning: mapear o territorio que cada concorrente "possui" na mente do consumidor social
5. Social-to-Revenue Attribution: conectar atividade social a impacto real em pipeline e receita

DOMINIOS DE EXPERTISE PROFUNDA:
- Metricas de Valor Real: Engagement Rate por Alcance (nao por followers), Save Rate (indicador de valor),
  Share Rate (indicador de advocacy), Comment Sentiment Score, Story Reply Rate, DM Conversion Rate
- Content Intelligence: performance por formato x tema x horario x tom; creative fatigue detection;
  trend-jacking strategy; meme marketing; UGC activation; employee advocacy programs
- Platform Mastery:
  * Instagram: Algorithm signals (saves>shares>comments>likes), Reels optimization, Collab posts,
    Broadcast Channels, Shopping integration, Creator partnerships
  * TikTok: FYP algorithm (watch time>completion rate>shares>comments), Spark Ads, TikTok Shop,
    Sound strategy, Duets/Stitches para engagement, Creator Marketplace
  * LinkedIn: SSI (Social Selling Index), Newsletter strategy, Document posts (3x alcance),
    Employee Advocacy ROI, LinkedIn Live, B2B thought leadership framework
  * Facebook: Groups strategy, Reels crossposting, Event marketing, Community building,
    Meta Business Suite optimization, Advantage+ creative
  * WhatsApp: Business API, Broadcast Lists, Communities, Click-to-WhatsApp Ads integration,
    Conversational Commerce, NPS via WhatsApp
- Social Commerce: Shoppable posts, Live Shopping, Social proof optimization, UGC-driven PDPs
- Social Listening Avancado: brand mention velocity, sentiment drift, crisis detection, competitive share of voice

FORMULAS QUE APLICA:
- True Engagement Rate = (Likes + Comments + Shares + Saves) / Alcance real x 100
- Content Efficiency Score = Total de engajamento / Numero de posts no periodo
- Share of Voice Social = Engajamento da marca / Engajamento total do setor x 100
- EMV (Earned Media Value) = Impressoes organicas x CPM equivalente do setor
- Social ROI = (Receita atribuida a social - Investimento em social) / Investimento x 100
- Community Health = (Advocacy rate x 0.4) + (Sentiment score x 0.3) + (Growth rate x 0.3)

BENCHMARKS QUE CITA COM AUTORIDADE (2024-2025, mercado Brasil):
- Instagram: ER 1.0-3.0% (bom), Save rate >2% (excelente), Reels completion >40% (bom)
- TikTok: ER 3.0-9.0%, watch time >50% (viral potential), share rate >1% (conteudo de valor)
- LinkedIn: ER 2.0-4.0%, SSI >70 (top performer), document posts 3x engagement vs texto puro
- Facebook: ER 0.06-0.15% (alcance organico quase morto), Reels ER 0.5-1.5% (unica esperanca)
- WhatsApp Business: open rate 98%, CTR 45-60%, response rate >80% para mensagens personalizadas
- Frequencia ideal: IG 5-7/semana, TT 1-3/dia, LI 3-5/semana, FB 1-2/dia (foco em Reels)""",

    'trafego': """VOCE E: O growth hacker mais experiente do Brasil. 30 anos otimizando funis de aquisicao,
de startups pre-seed a corporacoes com 1B+ de revenue. Ex-VP de Growth de unicornio brasileiro.
Certificado Google Analytics (todas as versoes desde Urchin), Adobe Analytics, Mixpanel, Amplitude.
Ja otimizou funis que geraram R$500M+ em receita incremental acumulada.

Voce NAO reporta metricas. Voce DIAGNOSTICA o ecossistema de aquisicao como um medico diagnostica um organismo --
identificando onde o sangue (trafego) flui bem, onde ha obstrucoes, e onde ha hemorragias de oportunidade.

METODOLOGIA PROPRIETARIA - MACFOR GROWTH INTELLIGENCE:
1. Traffic Ecosystem Health Check: diagnosticar a saude do ecossistema de aquisicao como um todo organico
2. Channel Dependency Risk Assessment: identificar vulnerabilidades de concentracao que ameacam o negocio
3. Cross-Channel Synergy Mapping: descobrir como canais se potencializam (ou se canibalizam) mutuamente
4. Acquisition Efficiency Frontier: plotar cada canal na fronteira eficiente (custo vs qualidade vs escala)
5. Growth Bottleneck Identification: encontrar O gargalo que mais limita o crescimento total

DOMINIOS DE EXPERTISE PROFUNDA:
- Atribuicao Multicanal: last-click (obsoleto), first-click, linear, time-decay, position-based,
  data-driven attribution (GA4), Marketing Mix Modeling (MMM), incrementality testing
- Funil AARRR Avancado: acquisition cost por canal, activation rate por cohort, retention curves (D1/D7/D30),
  revenue per user por fonte, referral coefficient (k-factor)
- Growth Loops: content loop (conteudo > SEO > trafego > mais conteudo), viral loop (usuario > convite > novo usuario),
  paid loop (investimento > receita > reinvestimento), sales loop (lead > venda > case > credibilidade > mais leads)
- CRO Avancado: multivariate testing, personalizacao por segmento, form optimization, page speed impact on conversion,
  cognitive load reduction, persuasion architecture (Cialdini), behavioral design

FORMULAS QUE APLICA:
- Traffic Quality Score = (Conv. Rate x Avg Session Duration x Pages/Session) / Bounce Rate
- Channel Dependency Index = Trafego do canal dominante / Trafego total x 100 (>50% = critico)
- Customer Acquisition Cost por canal = Investimento total no canal / Conversoes atribuidas
- Payback Period = CAC / (Revenue per customer por mes)
- Growth Rate necessario = (Meta de receita - Receita atual) / Receita atual / 12 meses
- Organic Asset Value = Trafego organico x CPC medio do setor x 12 = valor anual do ativo SEO

BENCHMARKS QUE CITA COM AUTORIDADE:
- Mix saudavel de trafego: Organico >35%, Direto 15-25%, Social 5-15%, Referral 5-10%, Pago <30%
- Dependencia de canal >50% = "Single Point of Failure" -- risco existencial para o negocio
- Bounce rate: B2B 25-55%, E-commerce 20-45%, Content/Blog 40-65%
- Conv. Rate: B2B SaaS 3-7%, E-commerce 1.5-3.5%, Lead Gen 5-12%
- Cada 1s de melhoria em page load = +7% em conversao (Deloitte Digital)
- Usuarios mobile: 74% do trafego web no Brasil (Statcounter) -- mobile-first e obrigatorio""",

    'midia_paga': """VOCE E: O estrategista de performance media mais respeitado do Brasil. 30 anos otimizando investimentos
em midia digital, desde os primeiros banners nos anos 90 ate Performance Max e AI-powered bidding em 2025.
Gestao acumulada de R$2B+ em budget de midia. Ex-diretor de midia da maior agencia de performance do Brasil.
Certificado Google Ads (todas as certificacoes), Meta Blueprint, LinkedIn Marketing Solutions, The Trade Desk.

Voce NAO gerencia campanhas. Voce ARQUITETA ecossistemas de midia que transformam investimento em receita
com eficiencia cirurgica. Cada real investido deve ter um destino estrategico claro.

METODOLOGIA PROPRIETARIA - MACFOR MEDIA INTELLIGENCE:
1. Investment Efficiency Audit: diagnosticar cada centavo investido vs retorno real (nao ROAS de plataforma)
2. Competitive Media Landscape: mapear investimento e estrategia de midia de cada concorrente
3. Budget Optimization Modeling: encontrar o ponto otimo de investimento por canal antes de diminishing returns
4. Paid-Organic Cannibalization Analysis: identificar quanto do investimento pago canibaliza trafego organico gratuito
5. Full-Funnel Attribution Reset: corrigir atribuicao para refletir o verdadeiro impacto de cada canal no revenue

DOMINIOS DE EXPERTISE PROFUNDA:
- Google Ads Avancado: Search (SKAGs, DSAs, RSAs optimization), Shopping (feed optimization, Merchant Center),
  Performance Max (asset group strategy, audience signals), YouTube (TrueView, Bumper, Shorts Ads),
  Demand Gen campaigns, bidding strategies avancadas (tROAS, tCPA, maximize conversion value)
- Meta Ads Avancado: Campaign Budget Optimization, Advantage+ Shopping, Advantage+ Creative,
  Conversions API (server-side tracking), Catalog Sales, Dynamic Ads, Lead Gen com Instant Forms,
  Lookalike audiences evolution, Broad targeting + creative testing methodology
- LinkedIn Ads: Sponsored Content, Message Ads, Conversation Ads, Document Ads, Event Ads,
  ABM (Account-Based Marketing) targeting, Matched Audiences, Revenue Attribution Reports
- Programatica: DSPs (DV360, The Trade Desk), DCO (Dynamic Creative Optimization), contextual targeting,
  viewability standards (MRC), brand safety (IAS, DoubleVerify), fraud detection, cross-device attribution
- Measurement & Attribution: GA4 data-driven attribution, Meta CAPI, Incrementality testing (geo-split, PSA),
  Marketing Mix Modeling, Conversion Lift Studies, Brand Lift Studies

FORMULAS QUE APLICA:
- True ROAS = Receita incremental atribuida / Investimento total (incluindo fees de agencia e tecnologia)
- Media Efficiency Ratio = Revenue / (Media Spend + Creative Cost + Tech Cost + Agency Fees)
- Incrementality Rate = (Conversoes com ads - Conversoes sem ads) / Conversoes com ads
- Keyword Waste Score = Investimento em keywords com Conv. Rate < 0.5% / Budget total
- Paid-Organic Overlap = Keywords onde rankeia top 3 organicamente E compra ads = desperdicio potencial
- LTV/CAC por canal: segmentar para encontrar canais que trazem clientes de MAIOR valor

BENCHMARKS QUE CITA COM AUTORIDADE (Brasil, 2024-2025):
- Google Ads: Search CPC R$1.50-4.00, CTR Search 3-8%, Conv Rate Search 2-5%
- Meta Ads: CPM R$15-40, CPC R$0.50-2.00, CTR 0.8-2.5%, CPL variavel (R$10-150 por industria)
- LinkedIn Ads: CPC R$8-25, CPM R$80-200, CTR 0.4-0.8%, CPL B2B R$50-300
- TikTok Ads: CPM R$8-20, CPC R$0.30-1.50, ideal para awareness e consideracao
- ROAS minimo: 3:1 e-commerce, 5:1 high-ticket B2C, 8:1+ para margem baixa
- LTV/CAC: >3:1 saudavel, 1-3:1 sustentavel com otimizacao, <1:1 INSUSTENTAVEL""",

    'buzz': """VOCE E: O maior decodificador de comportamento do consumidor digital do Brasil. 30 anos
transformando padroes de busca em mapas de oportunidade e tendencias de conteudo em vantagem competitiva.
Ex-Chief Content Officer do maior publisher digital da America Latina (50M+ visitors/mes).
Palestrante em Content Marketing World (Cleveland), Cannes Lions, SXSW Interactive.
Autor de metodologia de Content Intelligence adotada por 100+ marcas.

Voce NAO lista tendencias. Voce DECODIFICA o comportamento do consumidor como um antropologo digital,
revelando desejos, medos e intencoes ocultas por tras de cada busca no Google e cada scroll no feed.

METODOLOGIA PROPRIETARIA - MACFOR CONTENT INTELLIGENCE:
1. Consumer Search DNA: decodificar a jornada de busca completa do consumidor no segmento
2. Content White Space Mapping: identificar territorios de alto valor com baixa competicao
3. Demand Signal Analysis: usar padroes de busca como sinais preditivos de demanda de mercado
4. Cultural Moment Strategy: mapear momentos culturais que podem ser capturados pela marca
5. Competitive Content Moat Assessment: avaliar quao defensavel e a posicao de conteudo de cada player

DOMINIOS DE EXPERTISE PROFUNDA:
- Search Intent Mastery: decodificar intencao em 4 camadas (informational > commercial investigation >
  transactional > navigational) E sub-camadas (problem-aware > solution-aware > product-aware > brand-aware)
- Trend Intelligence: Google Trends (correlacao, sazonalidade, breakout), social listening trends,
  emerging topics detection, micro-trend vs macro-trend differentiation, trend lifecycle prediction
- Content Gap Analysis Avancada: nao apenas "keywords que faltam" mas territorios INTEIROS de conteudo
  que representam oportunidade de negocio (topicos x intencao x volume x competicao x ticket)
- Digital PR & Earned Media: data-driven storytelling que gera cobertura jornalistica espontanea,
  newsjacking framework (preparar > monitorar > reagir em <2h), expert commentary pipeline
- Consumer Psychology: principios de Cialdini aplicados a conteudo (reciprocidade, prova social,
  autoridade, escassez, comprometimento), neurocopywriting, behavioral content design
- SERP Feature Domination: PAA strategy (responder perguntas antes que o usuario as faca),
  featured snippet engineering, knowledge panel optimization, video carousel capture

FORMULAS QUE APLICA:
- Content Opportunity Score = (Volume de busca x CPC medio x Conv. Rate) / (Keyword Difficulty x Custo de producao)
- Topic Authority Index = Numero de keywords top 10 no cluster / Total de keywords do cluster x 100
- Content Decay Rate = (Trafego mes atual - Trafego pico) / Trafego pico x 100
- Share of Search = Volume de busca da marca / Volume total de buscas do setor x 100
  (Les Binet: Share of Search e o melhor preditor de market share futuro)
- Demand Elasticity = Variacao % no volume de busca / Variacao % no preco/oferta

REFERENCIAS QUE CITA COM AUTORIDADE:
- Google/Offerwise Brasil: 84% pesquisa online antes de comprar
- SparkToro/Datos: 58.5% das buscas sao zero-click (impacto em estrategia de conteudo)
- BrightEdge: 53% de todo trafego web vem de organic search
- Content Marketing Institute: empresas com blog geram 67% mais leads que sem
- HubSpot: conteudo composto (evergreen) gera 38% do trafego total de blogs maduros
- Backlinko: long-form content (3000+ palavras) recebe 77.2% mais backlinks que posts curtos
- Videos: 82% do trafego de internet em 2025 (Cisco), YouTube e o 2o maior buscador do mundo""",

    'aio': """VOCE E: O pioneiro ABSOLUTO em AI Search Optimization no Brasil e um dos 50 maiores especialistas
do mundo em GEO (Generative Engine Optimization). Pesquisador ativo desde o lancamento do ChatGPT (Nov 2022).
Publicou os primeiros estudos brasileiros sobre impacto de AI Overview no CTR organico.
Consultor de marcas Fortune 500 em estrategia de presenca em AI Search. Citado em Search Engine Journal,
Search Engine Land e Moz como referencia em GEO. Co-autor do primeiro framework de GEO Audit do mercado.

Voce esta na VANGUARDA de uma revolucao. Enquanto 95% do mercado ainda ignora AI Search, voce ja mapeou
o terreno e sabe exatamente como posicionar marcas para vencer na era da busca por IA.

METODOLOGIA PROPRIETARIA - MACFOR GEO INTELLIGENCE:
1. AI Citability Audit: avaliar a probabilidade da marca ser citada por ChatGPT, Gemini, Perplexity, Claude
2. SERP Disruption Forecast: prever quais keywords do setor serao mais impactadas por AI Overview
3. Entity Authority Mapping: mapear a presenca da marca no Knowledge Graph e em datasets de treinamento de LLMs
4. Content AI-Readiness Score: avaliar se o conteudo existente e "citavel" por IAs generativas
5. Defensive GEO Strategy: proteger trafego atual E capturar novas oportunidades em AI Search

DOMINIOS DE EXPERTISE PROFUNDA:
- AI Overview (Google): como funciona o sistema de selecao de fontes, quais signals pesam mais
  (authority, freshness, E-E-A-T, structured data, user intent alignment), como otimizar para ser citado
- LLM Citation Mechanics: como cada LLM decide quais fontes citar:
  * GPT-4/ChatGPT: prioriza fontes de alta autoridade, conteudo abrangente, reviews verificados, Wikipedia
  * Gemini: integrado com Google Search, prioriza fontes que ja rankeiam + Knowledge Graph + structured data
  * Perplexity: busca em tempo real, prioriza conteudo recente, bem citado, com dados originais e verificaveis
  * Claude: prioriza conteudo factual, bem estruturado, de fontes respeitaveis com autoria clara
- GEO (Generative Engine Optimization): diferente de SEO tradicional -- otimizar para ser A RESPOSTA,
  nao apenas um dos 10 links azuis. Foco em: clareza, autoridade, dados proprietarios, citabilidade
- Knowledge Graph Optimization: Google Knowledge Panel, Wikidata entities, schema markup avancado,
  entity disambiguation, brand entity strength assessment
- AI Brand Monitoring: monitorar como a marca e descrita/recomendada em respostas de ChatGPT, Gemini,
  Perplexity -- accuracy, sentiment, frequency, prominence, competitive positioning

FORMULAS QUE APLICA:
- AI Readiness Score = (Authority Score x 0.30) + (Content Quality E-E-A-T x 0.25) + (Structured Data x 0.20)
  + (External Mentions x 0.15) + (Brand Search Volume x 0.10) = Score 0-100
- AI Overview Risk = Keywords informacionais com AI Overview / Total de keywords organicas x 100
- CTR Impact Projection = Trafego atual em keywords de risco x Reducao media de CTR (25-40%) = Trafego perdido
- GEO Opportunity = Keywords sem AI Overview ainda + Keywords onde concorrente nao e citado = Janela de captura
- Citation Probability = f(DA, content quality, external mentions, schema, freshness, entity strength)

DADOS E TENDENCIAS QUE CITA COM AUTORIDADE:
- AI Overview ativado em 15-30% das buscas Google (variavel por vertical e pais, tendencia de alta)
- Impacto no CTR organico: -25% a -40% para queries informacionais, -10% a -20% para commercial investigation
- Perplexity: crescimento de 20x em 2024, 100M+ queries/mes, ameaca real ao modelo de busca tradicional
- SearchGPT/ChatGPT Search: 250M+ usuarios semanais do ChatGPT, integracao com busca em expansao
- Gen-Z: 40% prefere buscar no TikTok/ChatGPT vs Google (Adobe Survey 2024)
- Previsao Gartner: trafego organico tradicional caira 25% ate 2026 por causa de AI Search
- Conteudo com schema markup: 2-3x mais chances de citacao em AI Overview (estudo Zyppy/Rand Fishkin)
- Fatores de citacao em LLMs: autoridade (35%), conteudo (30%), mencoes externas (20%), structured data (15%)
- First-mover advantage em GEO: marcas que otimizam AGORA terao vantagem desproporcional quando AI Search for mainstream""",

    'estrategico': """VOCE E: Um CMO de elite / Senior Partner de estrategia digital. 30 anos no mais alto nivel
de decisao estrategica de marketing. Ex-CMO de 3 empresas listadas na B3 (Bovespa). Conselheiro de boards
de private equity e venture capital. MBA por Wharton + INSEAD. Especializado em digital transformation
e growth strategy. Liderou transformacoes digitais que geraram R$1B+ em valor de mercado.

Voce PENSA como um CEO, FALA como um McKinsey Senior Partner, e EXECUTA como o melhor CMO do mercado.
Seu diagnostico nao e uma opiniao -- e uma TESE DE INVESTIMENTO apoiada por dados.

METODOLOGIA PROPRIETARIA - MACFOR STRATEGIC INTELLIGENCE:
1. Digital Health Assessment: diagnostico de saude digital em 360 graus com scoring proprietario
2. Competitive Intelligence Matrix: mapa de posicao competitiva em CADA dimensao digital
3. Revenue Impact Modeling: traducao de CADA gap digital em impacto financeiro estimado
4. Strategic Opportunity Ranking: priorizacao cientifica (ICE Score) de TODAS as oportunidades identificadas
5. Transformation Roadmap: plano de execucao em ondas com OKRs, milestones e business case

DOMINIOS DE EXPERTISE PROFUNDA:
- Visao 360 Integrada: como SEO alimenta Content que alimenta Social que alimenta Paid que retroalimenta SEO
  -- voce ve o SISTEMA, nao canais isolados
- Business Translation: voce traduz "DA de 35" em "a marca tem 40% menos credibilidade digital que o lider,
  o que custa R$X/mes em trafego qualificado perdido e R$Y em CPC inflacionado"
- Financial Modeling: ROI por canal, LTV/CAC ratio, payback period, unit economics de aquisicao digital
- Scenario Planning: modelagem de 3 cenarios (otimista/base/pessimista) com impacto financeiro para cada
- Stakeholder Management: voce sabe que o CEO quer ouvir "receita e crescimento", o CFO quer "ROI e payback",
  e o board quer "vantagem competitiva sustentavel" -- e voce fala a lingua de cada um

FRAMEWORKS QUE APLICA:
- McKinsey Situation-Complication-Resolution para Executive Summaries impactantes
- Balanced Scorecard Digital adaptado: Financeiro, Cliente/Mercado, Processos/Canais, Inovacao/AI
- Matriz Esforco x Impacto com 4 quadrantes: Quick Wins, Projetos Estrategicos, Fill-ins, Thankless Tasks
- OKR Framework (Objectives & Key Results) para estruturar metas mensuráveis
- MECE Principle (McKinsey) para estruturacao logica sem sobreposicao e sem gaps
- 3 Horizons of Growth (McKinsey): H1 (core, 0-6m), H2 (adjacent, 6-12m), H3 (transformational, 12-24m)
- Digital Maturity Model: Nivel 1 (Ad hoc) > 2 (Oportunista) > 3 (Sistematico) > 4 (Estrategico) > 5 (Transformador)
- Porter's Value Chain Digital: onde a presenca digital cria (ou destroi) valor em cada elo
- Blue Ocean/Red Ocean: identificar se o prospect esta competindo em oceano vermelho ou pode criar oceano azul

ESTILO DE ENTREGA (inegociavel):
- Abre com "the killer headline" -- a descoberta que faz o CEO largar o celular e prestar atencao
- Cada insight: DADO (numero real) > SO WHAT (impacto no negocio) > NOW WHAT (acao) > QUANTO (projecao financeira)
- Nunca "achamos" ou "sugerimos" -- sempre "os dados demonstram" e "nossa recomendacao e"
- Fecha com roadmap de 12 meses com OKRs trimestrais e business case que convence o CFO
- Linguagem de boardroom: concisa, precisa, sem jargoes desnecessarios, orientada a decisao e investimento"""
}

def gerar_texto(prompt, sistema=SISTEMA_BASE, especialista=None):
    """Gera texto usando Anthropic Claude ou, como fallback, Gemini."""
    sistema_final = sistema
    if especialista and especialista in ESPECIALISTAS:
        sistema_final = f"{SISTEMA_BASE}\n\nPERFIL DO ESPECIALISTA:\n{ESPECIALISTAS[especialista]}"

    # Injeta contexto adicional do usuario se disponivel
    ctx = st.session_state.get('contexto_adicional', '')
    if ctx:
        sistema_final += f"\n\nCONTEXTO ADICIONAL FORNECIDO PELO ANALISTA DA AGENCIA (use para enriquecer e direcionar sua analise):\n{ctx}"

    if cliente_anthropic:
        try:
            response = cliente_anthropic.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=8096,
                system=sistema_final,
                messages=[{"role": "user", "content": prompt}]
            )
            return response.content[0].text
        except Exception:
            return gerar_texto_gemini(prompt, sistema_final)
    else:
        return gerar_texto_gemini(prompt, sistema_final)

def gerar_texto_gemini(prompt, sistema):
    """Fallback para gerar texto usando Gemini."""
    try:
        response = modelo_gemini.generate_content(f"{sistema}\n\n{prompt}")
        return response.text
    except Exception as e:
        return f"Erro ao gerar texto com Gemini: {str(e)}"

# =============================================================================
# FUNÇÕES DE CARREGAMENTO DE DADOS
# =============================================================================
def safe_float(val):
    try:
        return float(str(val).replace(',', '.').replace('R$', '').replace(' ', ''))
    except (ValueError, TypeError):
        return 0.0

def safe_int(val):
    try:
        return int(float(str(val).replace(',', '.').replace('.', '')))
    except (ValueError, TypeError):
        return 0

def carregar_csv(uploaded_file, separadores=[',', ';', '\t']):
    """Carrega um CSV para um DataFrame, tentando diferentes separadores e encodings."""
    if uploaded_file is not None:
        encodings = ['utf-8', 'latin1', 'cp1252', 'iso-8859-1']
        for enc in encodings:
            for sep in separadores:
                try:
                    uploaded_file.seek(0)
                    df = pd.read_csv(uploaded_file, sep=sep, encoding=enc, on_bad_lines='skip')
                    if df.shape[1] > 1:
                        # Limpa nomes de colunas
                        df.columns = [str(c).strip() for c in df.columns]
                        return df
                except:
                    continue

        # Ultima tentativa: ler como coluna unica e tentar auto-detect
        try:
            uploaded_file.seek(0)
            content = uploaded_file.read().decode('utf-8', errors='replace')
            df = pd.read_csv(StringIO(content), sep=None, engine='python', on_bad_lines='skip')
            if not df.empty:
                df.columns = [str(c).strip() for c in df.columns]
                return df
        except:
            pass
    return None

def carregar_e_combinar_csvs(arquivos):
    """Carrega multiplos CSVs e combina em um unico DataFrame."""
    if not arquivos:
        return None
    dfs = []
    for arq in arquivos:
        df = carregar_csv(arq)
        if df is not None and not df.empty:
            dfs.append(df)
    if not dfs:
        return None
    if len(dfs) == 1:
        return dfs[0]
    # Tenta concatenar; se colunas diferentes, concatena com outer join
    try:
        return pd.concat(dfs, ignore_index=True, sort=False)
    except Exception:
        return dfs[0]

def df_para_contexto(df, max_linhas=50):
    """Converte DataFrame em texto rico para enviar à IA, maximizando dados reais."""
    if df is None or (isinstance(df, pd.DataFrame) and df.empty):
        return ""
    if isinstance(df, dict):
        if not df:
            return ""
        try:
            df = pd.DataFrame(df)
        except Exception:
            return str(df)

    info = []
    info.append(f"Dimensoes: {df.shape[0]} linhas x {df.shape[1]} colunas")
    info.append(f"Colunas: {list(df.columns)}")

    # Estatísticas numéricas
    numericas = df.select_dtypes(include=['number'])
    if not numericas.empty:
        info.append(f"\nEstatisticas numericas:\n{numericas.describe().to_string()}")

    # Dados completos (até max_linhas)
    n = min(len(df), max_linhas)
    info.append(f"\nDados ({n} de {len(df)} linhas):\n{df.head(n).to_string()}")

    # Se tem mais linhas, mostra as últimas também para ver tendências
    if len(df) > max_linhas:
        info.append(f"\nUltimas 10 linhas:\n{df.tail(10).to_string()}")

    return "\n".join(info)


def interpretar_csv_com_ia(df, contexto_tipo):
    """Usa IA para interpretar a estrutura de um CSV arbitrário e extrair dados relevantes."""
    if df is None or df.empty:
        return ""

    prompt = f"""Analise a estrutura deste CSV e extraia os dados mais relevantes para uma análise de {contexto_tipo}.

{df_para_contexto(df, 30)}

INSTRUÇÕES:
1. Identifique o que cada coluna representa, mesmo que os nomes estejam abreviados ou em outro idioma
2. Extraia um resumo estruturado dos dados mais relevantes para análise de {contexto_tipo}
3. Identifique quais linhas/colunas representam o prospect vs concorrentes (se aplicável)
4. Destaque valores notáveis, outliers, tendências visíveis e comparativos entre players
5. Se houver dados temporais, identifique a direção das tendências (crescimento/queda)
6. Apresente os dados já interpretados, NÃO como tabela bruta mas como observações estruturadas

Formate a saída como um briefing analítico organizado por temas."""

    return gerar_texto(prompt, especialista=None)

# =============================================================================
# FUNÇÕES DE PROCESSAMENTO DE DADOS ESPECÍFICOS
# =============================================================================
def processar_dados_prospect_concorrentes(df):
    """Extrai nomes do prospect e concorrentes do CSV de cadastro."""
    prospect = ""
    concorrentes = []
    
    # Procura por linhas com "Prospect" ou "Nome"
    for idx, row in df.iterrows():
        row_str = ' '.join([str(x) for x in row.values if pd.notna(x)])
        if 'Prospect' in row_str and not prospect:
            # Tenta extrair o nome do prospect
            for val in row.values:
                if pd.notna(val) and 'Prospect' not in str(val) and len(str(val)) > 3:
                    prospect = str(val).strip()
                    break
        elif 'Concorrente' in row_str:
            for val in row.values:
                if pd.notna(val) and 'Concorrente' not in str(val) and len(str(val)) > 3 and 'Observações' not in str(val):
                    concorrentes.append(str(val).strip())
                    break
    
    # Se não encontrou, usa valores padrão
    if not prospect:
        prospect = "Prospect"
    if not concorrentes:
        concorrentes = ["Concorrente 1", "Concorrente 2", "Concorrente 3", "Concorrente 4"]
    
    return prospect, concorrentes[:4]  # Limita a 4 concorrentes

def processar_kw_principais(df):
    """Extrai as principais palavras-chave."""
    kws = []
    for idx, row in df.iterrows():
        row_str = ' '.join([str(x) for x in row.values if pd.notna(x)])
        if 'KW' in row_str or 'Keyword' in row_str:
            for val in row.values:
                if pd.notna(val) and 'KW' not in str(val) and 'Keyword' not in str(val) and len(str(val)) > 2:
                    kws.append(str(val).strip())
                    if len(kws) >= 10:
                        break
    return kws[:10]

def processar_seo_historico(df):
    """Processa dados históricos de SEO (rank, tráfego, keywords)."""
    # Identifica a estrutura: linhas com anos/meses, domínios, ranks, etc.
    dados = []
    
    # Procura por colunas que parecem datas (YYYY|MM)
    colunas_data = []
    for col in df.columns:
        if isinstance(col, str) and ('|' in col or re.match(r'\d{4}', col)):
            colunas_data.append(col)
    
    if colunas_data:
        # Encontra linhas com domínios
        for idx, row in df.iterrows():
            dominio = None
            for val in row.values[:2]:  # Primeiras colunas podem ter o domínio
                if pd.notna(val) and isinstance(val, str) and ('www.' in val or '.com' in val):
                    dominio = val
                    break
            
            if dominio:
                dados_row = {'dominio': dominio}
                for col in colunas_data:
                    if col in df.columns and idx < len(df):
                        dados_row[col] = row[col]
                dados.append(dados_row)
    
    return pd.DataFrame(dados) if dados else df

def processar_kw_ranking(df):
    """Processa dados de ranking de palavras-chave."""
    # Estrutura típica: Domain, Keyword, Position, Search volume, etc.
    colunas_esperadas = ['Domain', 'Keyword', 'Position', 'Search volume', 'URL']
    df_result = pd.DataFrame()
    
    for col in colunas_esperadas:
        for df_col in df.columns:
            if col.lower() in str(df_col).lower():
                df_result[col] = df[df_col]
                break
    
    return df_result if not df_result.empty else df

def processar_analise_kw(df):
    """Processa análise de palavras-chave com volumes e CPC."""
    # Procura por colunas: Keyword, Search volume, CPC, Competition
    colunas_mapeadas = {}
    mapeamento = {
        'keyword': 'Keyword',
        'search volume': 'Search volume',
        'cpc': 'CPC',
        'competition': 'Competition',
        'number of results': 'Number of results'
    }
    
    for termo, nome in mapeamento.items():
        for col in df.columns:
            if termo.lower() in str(col).lower():
                colunas_mapeadas[nome] = col
                break
    
    if colunas_mapeadas:
        df_result = df[list(colunas_mapeadas.values())].copy()
        df_result.columns = list(colunas_mapeadas.keys())
        return df_result
    
    return df

def processar_social_facebook(df):
    """Processa dados de engajamento do Facebook."""
    dados = {}
    
    # Procura por linhas com o nome do prospect
    for idx, row in df.iterrows():
        row_str = ' '.join([str(x) for x in row.values if pd.notna(x)])
        if 'Prospect' in row_str or 'plumatex' in row_str.lower():
            # Extrai métricas
            for j, val in enumerate(row.values):
                if pd.notna(val) and isinstance(val, (int, float)) or (isinstance(val, str) and val.replace('.', '').replace(',', '').isdigit()):
                    if 'likes' in str(row.index[j]).lower() or 'followers' in str(row.index[j]).lower():
                        dados['likes'] = safe_int(val)
                    elif 'engagement' in str(row.index[j]).lower():
                        dados['engagement'] = safe_float(val)
                    elif 'posts' in str(row.index[j]).lower():
                        dados['posts'] = safe_int(val)
    
    return dados

def processar_social_instagram(df):
    """Processa dados do Instagram."""
    dados = {}
    
    for idx, row in df.iterrows():
        row_str = ' '.join([str(x) for x in row.values if pd.notna(x)])
        if 'Prospect' in row_str or 'plumatex' in row_str.lower():
            for j, val in enumerate(row.values):
                if pd.notna(val):
                    if 'followers' in str(row.index[j]).lower():
                        dados['followers'] = safe_int(val)
                    elif 'engagement' in str(row.index[j]).lower():
                        dados['engagement'] = safe_float(val)
    
    return dados

def processar_autoridade(df):
    """Processa dados de autoridade de domínio."""
    dados = {}
    
    for idx, row in df.iterrows():
        row_str = ' '.join([str(x) for x in row.values if pd.notna(x)])
        if 'plumatex' in row_str.lower() or (idx == 0 and 'Authority' in str(row.values)):
            for j, val in enumerate(row.values):
                if pd.notna(val) and isinstance(val, (int, float)) or (isinstance(val, str) and val.replace('.', '').isdigit()):
                    col_name = str(row.index[j]) if j < len(row.index) else f"col{j}"
                    if 'authority' in col_name.lower() or 'score' in col_name.lower():
                        dados['authority_score'] = safe_int(val)
                    elif 'domains' in col_name.lower():
                        dados['domains'] = safe_int(val)
                    elif 'follow' in col_name.lower():
                        dados['follow_links'] = safe_int(val)
    
    return dados

# =============================================================================
# FUNCOES DE GERACAO DE INSIGHTS COM IA
# =============================================================================

REGRAS_ANALISE = """
=== REGRAS INEGOCIAVEIS DE ANALISE ===

1. USE EXCLUSIVAMENTE OS DADOS REAIS FORNECIDOS. Cada numero nos CSVs e real e foi extraido de ferramentas
   profissionais (SEMrush, Ahrefs, SimilarWeb, etc). Cite valores exatos. Nao arredonde sem necessidade.

2. NUNCA INVENTE DADOS. Se uma metrica nao esta nos CSVs, diga "dado nao disponivel nos dados fornecidos"
   e siga para o proximo ponto. NAO ESTIME, NAO PROJETE, NAO CRIE benchmarks fictícios.

3. CRUZAMENTO OBRIGATORIO. Compare metricas entre players nos MESMOS periodos.
   Calcule diferencas percentuais, gaps absolutos, taxas de variacao.

4. ZERO GENERALIDADES. Nao escreva "o trafego e baixo". Escreva:
   "O trafego organico de [prospect] (X visitas/mes) representa Y% do lider [concorrente] (Z visitas/mes),
   um gap de W visitas qualificadas perdidas mensalmente."

5. CADA OBSERVACAO = DADO REAL + "E DAI?" (impacto no negocio) + ACAO RECOMENDADA.
   Incorpore os insights naturalmente na analise, nao como secao separada.

6. TENDENCIAS TEMPORAIS: Se houver dados de multiplos periodos, calcule e destaque:
   - Taxas de crescimento/queda mes a mes ou ano a ano
   - Pontos de inflexao (mudancas bruscas) e possiveis causas
   - Aceleracao ou desaceleracao de tendencias

7. QUANDO DADOS FOREM LIMITADOS para determinada metrica, aprofunde-se MAIS nos dados que ESTAO disponiveis.
   Nunca compense falta de dados com generalizacoes.
"""

def gerar_insights_seo(dados_brutos_texto, kw_principais, prospect, concorrentes):
    """Gera analise de SEO com dados reais."""

    conc_str = ', '.join(concorrentes) if concorrentes else 'identificar nos dados'
    kw_str = ', '.join(kw_principais) if kw_principais else 'extrair dos dados'

    prompt = f"""{REGRAS_ANALISE}

PROSPECT: {prospect}
CONCORRENTES: {conc_str}
PALAVRAS-CHAVE DO SEGMENTO: {kw_str}

DADOS REAIS EXTRAIDOS DOS CSVs:
{dados_brutos_texto}

Produza uma analise estrategica de SEO de nivel senior consultant. Use os frameworks e benchmarks abaixo
para contextualizar CADA dado encontrado nos CSVs.

## 1. SCORECARD COMPETITIVO DE SEO
Monte uma tabela comparativa com VALORES REAIS de cada player encontrado nos dados:
| Metrica | {prospect} | Concorrente 1 | ... | Lider | Gap vs Lider |
Metricas: Organic Traffic, Traffic Cost, Organic Keywords, Authority Score, Backlinks, Referring Domains.
Para CADA metrica, calcule:
- Gap absoluto e percentual vs lider do mercado
- Traffic Cost Equivalent: valor monetario do trafego organico (trafego x CPC medio do setor)
- Share of Voice organico: % das keywords do setor onde cada player aparece no top 10

## 2. INTELIGENCIA DE KEYWORDS (ANALISE PROFUNDA)
Analise CADA keyword real dos dados aplicando o modelo See-Think-Do-Care:
- SEE (awareness): keywords informacionais de topo de funil - volume alto, CPC baixo
- THINK (consideration): keywords de comparacao, "melhor", "vs" - volume medio, CPC medio
- DO (conversion): keywords transacionais, "comprar", "preco", "onde" - volume baixo, CPC alto
- CARE (retention): keywords de suporte, "como usar", "assistencia"

Para cada keyword calcule o VALOR ECONOMICO:
  Valor/mes = Volume x CTR da posicao atual (Backlinko: #1=31.7%, #2=24.7%, #3=18.7%, #5=9.5%, #10=3.1%)
              x Taxa de conversao media do setor x Ticket medio estimado

Identifique:
- Keyword Gaps criticos: termos de alto valor transacional onde {prospect} NAO aparece no top 20
- Quick Wins: keywords nas posicoes 11-20 que podem subir ao top 10 com otimizacao on-page
- Keywords Canibalizadas: multiplas URLs competindo pela mesma keyword
- Oportunidades Long-tail: clusters de keywords de cauda longa com intencao de compra clara

## 3. ANALISE DE AUTORIDADE E LINK PROFILE
Usando os dados de Authority Score, backlinks e referring domains:
- Classifique cada player: 0-20 (iniciante), 21-40 (em desenvolvimento), 41-60 (intermediario), 61-80 (forte), 80+ (dominante)
- Ratio backlinks/referring domains: alto ratio = poucos sites com muitos links (risco), baixo = perfil diversificado (saudavel)
- Velocidade de aquisicao de links: crescimento mes a mes dos referring domains
- Estimativa de investimento necessario em link building para fechar o gap (baseado no custo medio de R$500-2000 por link de qualidade)

## 4. EVOLUCAO TEMPORAL E TENDENCIAS
Se houver dados de multiplos periodos:
- CAGR (Compound Annual Growth Rate) de trafego organico para cada player
- Pontos de inflexao e correlacao com Google Core Updates conhecidos (Nov 2024, Mar 2024, Oct 2023, etc.)
- Analise de momentum: quem esta acelerando vs desacelerando
- Projecao de cenarios: otimista (best practices implementadas) / base (status quo) / pessimista (concorrentes aceleram)

## 5. OPORTUNIDADES DE SERP FEATURES
Com base nas keywords:
- Featured Snippets: quais keywords tem snippet e quem domina
- People Also Ask: perguntas relacionadas que podem gerar trafego incremental
- Local Pack: se relevante para o segmento
- Video Carousel: oportunidades de video SEO

## 6. CONCLUSAO ESTRATEGICA (FORMATO EXECUTIVO)
**Headline**: A descoberta mais impactante em uma frase (com numeros)
**Diagnostico**: Nivel de maturidade SEO de {prospect} (escala 1-5) com justificativa baseada nos dados
**Top 3 Gaps Criticos**: Cada um com valor economico estimado do que esta sendo perdido
**Roadmap**:
- Quick Wins (0-30 dias): acoes de otimizacao on-page, technical fixes -- resultado esperado: +X% trafego
- Projetos Taticos (1-3 meses): content gaps, link building inicial -- resultado esperado: +Y% keywords top 10
- Estrategia de Longo Prazo (3-12 meses): topic authority, digital PR -- resultado esperado: +Z% share of voice
**Business Case**: Investimento estimado vs retorno projetado em trafego, leads e receita"""

    return gerar_texto(prompt, especialista='seo')


def gerar_insights_social(dados_brutos_texto, prospect, concorrentes):
    """Gera analise de Social Media com dados reais."""

    conc_str = ', '.join(concorrentes) if concorrentes else 'identificar nos dados'

    prompt = f"""{REGRAS_ANALISE}

PROSPECT: {prospect}
CONCORRENTES: {conc_str}

DADOS REAIS EXTRAIDOS DOS CSVs:
{dados_brutos_texto}

Produza uma analise estrategica de Social Media de nivel VP/Head. Use os frameworks e benchmarks abaixo.

## 1. SCORECARD DE SOCIAL MEDIA
Monte tabela comparativa por plataforma (Facebook, Instagram, TikTok, LinkedIn, WhatsApp) com VALORES REAIS:
| Metrica | {prospect} | Concorrentes | Benchmark Setor | Avaliacao |
Metricas por plataforma:
- Followers/Base, Engagement Rate, Posts no periodo, Engajamento por post, Crescimento de base
- Para cada metrica, classifique: Abaixo do benchmark / Na media / Acima do benchmark

BENCHMARKS DE REFERENCIA (use para contextualizar):
- Facebook: ER medio 0.06-0.15%, alcance organico 2-5% da base
- Instagram: ER medio 1.0-3.0%, Reels 2-3x alcance do feed
- TikTok: ER medio 3.0-9.0%, completion rate >50% = bom
- LinkedIn: ER medio 2.0-4.0%, document posts 3x mais alcance
- Engagement Rate = (Curtidas + Comentarios + Compartilhamentos + Saves) / Alcance x 100

## 2. SOCIAL MEDIA MATURITY ASSESSMENT
Classifique {prospect} no Social Media Maturity Model:
- Nivel 1 (Presenca): perfis criados, publicacoes esporadicas
- Nivel 2 (Engajamento): frequencia consistente, conteudo planejado
- Nivel 3 (Comunidade): interacao ativa, UGC, community management
- Nivel 4 (Advocacy): embaixadores da marca, social proof, viralizacao organica
- Nivel 5 (Revenue): social commerce ativo, atribuicao clara social > venda

## 3. CONTENT PERFORMANCE MATRIX
Classifique os tipos de conteudo (se dados disponiveis) na matriz BCG adaptada:
- STARS: alto alcance + alto engajamento (escalar investimento)
- CASH COWS: engajamento consistente, alcance moderado (manter frequencia)
- QUESTION MARKS: alto alcance, baixo engajamento (testar variacoes)
- DOGS: baixo alcance e engajamento (descontinuar ou pivotar)

Analise por formato: video vs imagem vs carrossel vs stories vs reels vs texto puro
Analise por tema: quais territorios de conteudo performam melhor e por que

## 4. SHARE OF VOICE SOCIAL
Calcule o Share of Voice estimado:
- SOV = Engajamento total de {prospect} / Engajamento total de todos os players analisados x 100
- Correlacao SOV vs Share of Market (Les Binet & Peter Field): marcas com SOV > SOM tendem a crescer
- Excess Share of Voice (ESOV) = SOV - SOM: se positivo, indica potencial de crescimento

## 5. ANALISE COMPETITIVA PROFUNDA
Para CADA concorrente vs {prospect}:
- Estrategia de conteudo aparente: pilares tematicos, tom, frequencia, formatos
- Pontos fortes a observar e aprender
- Vulnerabilidades explorveis: gaps de conteudo, baixo engajamento em areas especificas
- Estimativa de investimento: baseado na frequencia e qualidade, estimativa do tamanho da equipe/budget

## 6. CONCLUSAO ESTRATEGICA
**Headline**: Descoberta principal (com numeros)
**Diagnostico de Maturidade**: nivel atual e nivel alvo em 12 meses
**Brand Health Score Social**: avaliacao de 0-100 baseada nos dados
**Top 3 Gaps**: com impacto estimado em brand awareness e pipeline
**Roadmap Social Media**:
- Imediato (0-30 dias): otimizacoes de perfil, ajuste de frequencia, formatos prioritarios
- Curto Prazo (1-3 meses): nova estrategia de conteudo, testes de formato, community building
- Medio Prazo (3-6 meses): influencer strategy, social commerce, UGC programs
**Metricas de Sucesso**: KPIs mensuraveis para cada fase"""

    return gerar_texto(prompt, especialista='social')


def gerar_insights_trafego(dados_brutos_texto, prospect, concorrentes):
    """Gera analise de fontes de trafego com dados reais."""

    conc_str = ', '.join(concorrentes) if concorrentes else 'identificar nos dados'

    prompt = f"""{REGRAS_ANALISE}

PROSPECT: {prospect}
CONCORRENTES: {conc_str}

DADOS REAIS EXTRAIDOS DOS CSVs:
{dados_brutos_texto}

Produza uma analise estrategica de Fontes de Trafego usando o RACE Framework e Pirate Metrics.

## 1. DIAGNOSTICO DO ECOSSISTEMA DE AQUISICAO
Para CADA player encontrado nos dados, identifique:
- Distribuicao por canal: organico, pago, direto, social, referral, email, display
- Channel Dependency Index: % do trafego que vem do canal dominante
  (>50% de um canal = RISCO CRITICO, 30-50% = risco moderado, <30% = saudavel)
- Motor principal de crescimento: qual canal mais contribui para trafego qualificado

BENCHMARKS DE MIX SAUDAVEL:
- Organico: >30% (indica maturidade SEO e brand equity)
- Direto: 15-25% (indica forca de marca)
- Social: 5-15% (indica engajamento de comunidade)
- Referral: 5-10% (indica autoridade e parcerias)
- Pago: <30% (proporcao maior indica dependencia de investimento)

## 2. ANALISE DE EFICIENCIA POR CANAL (RACE Framework)
Para cada fonte de trafego, avalie no framework RACE:
- REACH: volume de trafego gerado por canal
- ACT: engajamento (paginas/sessao, duracao, bounce rate por fonte)
- CONVERT: taxa de conversao por fonte (se dados disponiveis)
- ENGAGE: retorno/recorrencia por fonte

Calcule o Traffic Cost Equivalent:
- Trafego organico x CPC medio do setor = valor do ativo SEO
- Trafego social x custo equivalente de ads sociais = valor do social organico
- Trafego direto = proxy de brand equity digital (valuation de marca)

## 3. ANALISE DE VULNERABILIDADE
Aplique stress test ao mix de aquisicao:
- "E se o Google penalizar o site?" -- quanto trafego seria perdido? (cenario -50% organico)
- "E se cortar budget de ads?" -- quanto trafego cai? (cenario -100% pago)
- "E se algoritmo social mudar?" -- impacto no trafego social?
Para cada cenario, calcule: perda de trafego > perda estimada de leads > impacto em receita

## 4. TENDENCIAS TEMPORAIS E MOMENTUM
Se houver dados de multiplos periodos:
- CAGR por canal: qual esta crescendo mais rapido
- Sazonalidade: picos e vales recorrentes
- Correlacao entre canais: investimento em ads aumenta trafego direto? (halo effect)
- Ponto de inflexao: mudancas bruscas e possiveis causas

## 5. CONCLUSAO ESTRATEGICA
**Headline**: Vulnerabilidade ou oportunidade principal (com numeros)
**Health Score do Ecossistema**: nota 0-100 baseada em diversificacao, crescimento e qualidade
**Top 3 Riscos**: com probabilidade e impacto estimado
**Roadmap de Diversificacao**:
- Curto Prazo: rebalancear mix com quick wins por canal
- Medio Prazo: desenvolver canais subdimensionados
- Longo Prazo: construir moats (vantagens competitivas sustentaveis) em canais proprios"""

    return gerar_texto(prompt, especialista='trafego')


def gerar_insights_midia_paga(dados_brutos_texto, prospect, concorrentes):
    """Gera analise de midia paga com dados reais."""

    conc_str = ', '.join(concorrentes) if concorrentes else 'identificar nos dados'

    prompt = f"""{REGRAS_ANALISE}

PROSPECT: {prospect}
CONCORRENTES: {conc_str}

DADOS REAIS EXTRAIDOS DOS CSVs:
{dados_brutos_texto}

Produza uma analise estrategica de Midia Paga de nivel Head de Performance. Use frameworks de eficiencia
e benchmarks do mercado brasileiro.

## 1. SCORECARD DE INVESTIMENTO COMPETITIVO
Monte tabela comparativa com VALORES REAIS de cada player:
| Metrica | {prospect} | Concorrentes | Benchmark | Avaliacao |
Metricas: Paid Traffic, Paid Cost, CPC medio, % trafego pago vs total, Keywords compradas

Para cada player calcule:
- Custo por visita paga = Paid Cost / Paid Traffic
- Eficiencia relativa: quem extrai mais trafego por real investido
- Share of Spend estimado: % do investimento total do setor

BENCHMARKS DE REFERENCIA (mercado Brasil):
- Google Ads CPC medio: Search R$1.50-4.00, Display R$0.30-0.80, Shopping R$0.50-1.50
- Meta Ads: CPM R$15-40, CPC R$0.50-2.00, CPL variavel por industria
- LinkedIn Ads: CPC R$8-25 (B2B premium)
- ROAS minimo saudavel: 3:1 e-commerce, 5:1 high-ticket, 2:1 awareness

## 2. ANALISE DE EFICIENCIA (FRAMEWORK FULL-FUNNEL)
Avalie o investimento em cada etapa do funil:
- TOFU (Awareness): CPM, Reach, Video Views -- investimento em construcao de demanda
- MOFU (Consideration): CPC, CTR, Engagement -- investimento em educacao e consideracao
- BOFU (Conversion): CPA, ROAS, Conv. Rate -- investimento em captura de demanda
Se dados forem limitados, analise o que os dados de PPC revelam sobre a estrategia de cada player.

## 3. ANALISE DE DESPERDICIO E OPORTUNIDADE
- Sobreposicao pago-organico: keywords onde ja rankeia top 3 organicamente mas paga ads = DESPERDICIO
  Calcule: CPC dessas keywords x clicks estimados = valor desperdicado/mes
- Oportunidades de arbitragem: keywords baratas (CPC baixo) com alta intencao transacional
- Competitive gap: keywords que concorrentes compram e {prospect} nao (e vice-versa)
- Budget allocation: proporcao search vs display vs social ads -- esta otimizado?

## 4. ESTIMATIVA DE ROI E CENARIOS
Modele 3 cenarios de investimento:
- Conservador: otimizar budget atual com melhor alocacao -- impacto estimado
- Moderado: +30% budget com foco nas oportunidades identificadas -- impacto estimado
- Agressivo: dobrar investimento com full-funnel strategy -- impacto estimado
Para cada: investimento > trafego esperado > leads estimados > receita projetada (usando conv. rates do setor)

## 5. CONCLUSAO ESTRATEGICA
**Headline**: Principal achado sobre eficiencia/ineficiencia de investimento
**Efficiency Score**: nota 0-100 baseada em CPC, ROAS, mix de canais
**Top 3 Oportunidades**: com ROI projetado para cada uma
**Roadmap de Otimizacao**:
- Imediato: eliminar desperdicios, pausar keywords ineficientes
- Curto Prazo: redistribuir budget para canais/keywords de melhor performance
- Medio Prazo: expandir para novas plataformas e formatos com budget incremental"""

    return gerar_texto(prompt, especialista='midia_paga')


def gerar_insights_buzz(prospect, kw_principais=None):
    """Gera analise de buzz marketing e comportamento de busca."""

    kws = ', '.join(kw_principais) if kw_principais else 'nao fornecidas -- inferir do segmento do prospect'

    prompt = f"""{REGRAS_ANALISE}

PROSPECT: {prospect}
PALAVRAS-CHAVE DO SEGMENTO: {kws}

NOTA ESPECIAL: Para Content Intelligence, combine dados dos CSVs (se houver) com seu profundo
conhecimento do mercado digital brasileiro. Cada afirmacao deve ser HIPER-ESPECIFICA ao segmento
de {prospect}. Nada generico. Use o framework Content-Market Fit e Search Demand Curve.

Produza uma analise de Content Intelligence de nivel director.

## 1. MAPA DE INTENCAO DE BUSCA DO CONSUMIDOR
Usando o modelo See-Think-Do-Care e Jobs-To-Be-Done, mapeie o comportamento de busca:

**Jornada de Busca do Consumidor no segmento de {prospect}:**
- DESCOBERTA: queries informacionais de topo de funil (volume estimado, tendencia)
  Liste 10+ queries reais que consumidores fazem, agrupadas por tema
- CONSIDERACAO: queries comparativas, "melhor", "review", "vale a pena" (volume, tendencia)
  Liste 8+ queries de comparacao especificas do segmento
- DECISAO: queries transacionais, "comprar", "preco", "onde encontrar" (volume, CPC estimado)
  Liste 5+ queries de conversao de alto valor
- POS-COMPRA: "como usar", "assistencia", "troca" -- oportunidades de retencao

Para cada grupo: volume estimado de buscas mensais, tendencia (crescendo/estavel/caindo),
e o que revela sobre dores e necessidades nao atendidas.

## 2. ANALISE DE DEMANDA E SAZONALIDADE
Especifico para o segmento de {prospect}:
- Picos de demanda: quais meses/periodos tem mais buscas e por que
- Eventos gatilho: datas comerciais, estacoes, eventos setoriais que impulsionam demanda
- Micro-tendencias emergentes: temas em ascensao que ainda nao saturaram
- Macro-tendencias: mudancas estruturais no comportamento do consumidor do segmento

## 3. TERRITORIOS DE CONTEUDO (CONTENT-MARKET FIT)
Aplique a Search Demand Curve para identificar oportunidades:
- HEAD TERMS (alto volume, alta competicao): quais {prospect} PRECISA disputar para credibilidade
- MID-TAIL (volume medio, competicao moderada): oportunidades realistas de rankear em 3-6 meses
- LONG-TAIL (baixo volume, alta conversao): nichos de conteudo com compradores qualificados

Para cada territorio, avalie:
- Dificuldade de competicao (quem ja domina e qual a autoridade deles)
- Content gap: alta demanda + baixa qualidade de conteudo existente = OPORTUNIDADE
- Formato ideal: artigo longo, video, ferramenta, calculadora, comparativo, guia

## 4. ANALISE DE SENTIMENTO E REPUTACAO DIGITAL
Baseado no conhecimento do segmento:
- Onde consumidores falam sobre marcas do segmento (Reclame Aqui, Google Reviews, Reddit, TikTok, forums)
- Temas de insatisfacao recorrentes no setor (oportunidade de se diferenciar)
- Nivel de confianca digital do setor: o consumidor confia em comprar online?
- User-generated content: o que consumidores postam sobre o segmento espontaneamente

## 5. ESTRATEGIA DE CONTEUDO PESO (Paid, Earned, Shared, Owned)
Recomende mix de distribuicao para {prospect}:
- OWNED: blog, site, newsletter -- pilares de conteudo proprietario
- EARNED: digital PR, mencoes espontaneas, backlinks editoriais -- como conquistar
- SHARED: social media, UGC, parcerias de conteudo -- como amplificar
- PAID: content amplification, native ads, sponsored content -- como acelerar

## 6. CONCLUSAO ESTRATEGICA
**Headline**: A maior oportunidade de conteudo para {prospect} (especifica, com estimativa de volume)
**Content Maturity Score**: avaliacao 0-100 da maturidade de conteudo atual
**Top 5 Territorios Prioritarios**: rankeados por potencial de trafego x dificuldade x valor de negocio
**Calendario Editorial Estrategico**: temas prioritarios por trimestre, alinhados a sazonalidade
**Metricas de Sucesso**: KPIs por territorio (trafego, rankings, leads, engagement)"""

    return gerar_texto(prompt, especialista='buzz')


def gerar_insights_aio(dados_brutos_texto, prospect, kw_principais=None):
    """Gera analise de AI Search / GEO."""

    kws = ', '.join(kw_principais) if kw_principais else 'inferir do segmento do prospect'

    prompt = f"""{REGRAS_ANALISE}

PROSPECT: {prospect}
PALAVRAS-CHAVE DO SEGMENTO: {kws}

DADOS REAIS EXTRAIDOS DOS CSVs (metricas de authority, keywords, trafego relevantes para GEO):
{dados_brutos_texto}

Produza uma analise PIONEIRA de AI Search Optimization (GEO) -- este e um diferencial competitivo
da Macfor que pouquissimas agencias oferecem. A analise deve ser visionaria e pratica ao mesmo tempo.

## 1. AI SEARCH READINESS SCORE
Avalie a preparacao de {prospect} para o novo paradigma de busca por IA.
Calcule um score de 0-100 baseado em:
- Authority Score / Domain Rating (peso 35%): autoridade e confiabilidade percebida por LLMs
  Referencia: DA <30 = muito baixo para citacao, 30-50 = possivel, 50-70 = provavel, 70+ = alta probabilidade
- Qualidade de Conteudo E-E-A-T (peso 30%): expertise, experiencia, autoridade, confianca
  Analise: o conteudo do prospect demonstra autoria especialista, fontes primarias, dados originais?
- Dados Estruturados (peso 20%): schema markup, FAQs estruturadas, knowledge graph signals
- Mencoes Externas e Citacoes (peso 15%): backlinks de fontes autoritativas, mencoes em midia, reviews

## 2. IMPACTO DO AI OVERVIEW NAS KEYWORDS DO SETOR
Classifique as keywords dos dados em niveis de risco de AI Overview:
- ALTO RISCO (provavel AI Overview): queries informacionais, "como", "o que e", "melhor X para Y"
  Impacto estimado no CTR: -25 a -40% vs busca tradicional
- MEDIO RISCO (possivel AI Overview): queries comparativas, "X vs Y", reviews
  Impacto estimado no CTR: -10 a -25%
- BAIXO RISCO (improvavel AI Overview): queries transacionais diretas, navegacionais, locais
  Impacto minimo no CTR

Para keywords de alto risco, calcule:
- Trafego atual nessas keywords x reducao estimada de CTR = PERDA PROJETADA de trafego
- Quem esta posicionado para SER CITADO no AI Overview (authority + conteudo + entidade no knowledge graph)

## 3. ANALISE DE PRESENCA EM LLMs
Avalie como {prospect} provavelmente aparece (ou nao) em respostas de:
- ChatGPT/GPT-4: prioriza fontes com alta autoridade, conteudo abrangente, reviews verificados
- Google Gemini/AI Overview: prioriza fontes que ja rankeiam bem + dados estruturados + E-E-A-T
- Perplexity: prioriza conteudo recente, bem citado, com dados originais
- Claude: prioriza conteudo factual, bem estruturado, de fontes respeitadas

Fatores que determinam citacao em LLMs (pesquisa de GEO):
- Autoridade do dominio: 35% do peso
- Qualidade e profundidade do conteudo: 30%
- Mencoes externas e citacoes: 20%
- Dados estruturados e markup: 15%

## 4. ESTRATEGIA GEO (GENERATIVE ENGINE OPTIMIZATION)
Acoes concretas priorizadas por impacto:

**Otimizacao de Conteudo para LLMs:**
- Estrutura de conteudo: perguntas claras como headers, respostas diretas no primeiro paragrafo
- Dados proprietarios: pesquisas originais, benchmarks unicos, estudos de caso -- LLMs preferem dados exclusivos
- Autoria especialista: bylines de especialistas reais, credenciais visiveis, links para perfis
- Atualizacao continua: conteudo com data de atualizacao recente tem prioridade

**Schema Markup Prioritario:**
- FAQ Schema em todas as paginas de conteudo informacional
- HowTo Schema para tutoriais e guias
- Product Schema com reviews agregados
- Organization Schema com founding date, awards, credentials
- Article Schema com author, datePublished, dateModified

**Estrategia de Entidade e Knowledge Graph:**
- Criar/otimizar Google Knowledge Panel da marca
- Wikipedia/Wikidata presence (se elegivel)
- Consistencia de NAP (Name, Address, Phone) em todas as plataformas
- Citacoes em fontes autoritativas do setor

## 5. CENARIOS FUTUROS (2025-2027)
Modelagem de impacto da evolucao do AI Search no setor de {prospect}:
- Cenario 1 (conservador): AI Overview em 20% das buscas do setor, CTR organico cai 10-15%
- Cenario 2 (moderado): AI Overview em 40% das buscas, CTR cai 20-30%, novos canais de IA emergem
- Cenario 3 (agressivo): 60%+ das buscas mediadas por IA, modelo de trafego muda fundamentalmente
Para cada cenario: impacto em trafego, leads e receita de {prospect}

## 6. CONCLUSAO ESTRATEGICA
**Headline**: Nivel de preparacao de {prospect} para o futuro da busca
**AI Readiness Score**: X/100 com breakdown por fator
**Risco Quantificado**: % do trafego atual vulneravel a AI Overview
**Roadmap GEO**:
- Imediato (0-30 dias): schema markup, otimizacao de headers FAQ, autoria especialista
- Curto Prazo (1-3 meses): conteudo data-driven, pesquisas originais, link building autoritativo
- Medio Prazo (3-6 meses): knowledge graph optimization, presenca em plataformas de IA
- Longo Prazo (6-12 meses): AI-first content strategy, ferramentas interativas, dados proprietarios"""

    return gerar_texto(prompt, especialista='aio')


def gerar_recomendacoes_estrategicas(insights_seo, insights_social, insights_trafego, insights_midia, insights_buzz, insights_aio, prospect):
    """Consolida recomendacoes em plano estrategico."""

    # Monta material disponivel (so inclui o que tem conteudo)
    material = ""
    if insights_seo:
        material += f"\n--- SEO (Head de SEO, 30 anos exp.) ---\n{insights_seo[:4000]}\n"
    if insights_social:
        material += f"\n--- SOCIAL MEDIA (Head de Social, 30 anos exp.) ---\n{insights_social[:4000]}\n"
    if insights_trafego:
        material += f"\n--- FONTES DE TRAFEGO (Head de Growth, 30 anos exp.) ---\n{insights_trafego[:3000]}\n"
    if insights_midia:
        material += f"\n--- MIDIA PAGA (Head de Performance, 30 anos exp.) ---\n{insights_midia[:3000]}\n"
    if insights_buzz:
        material += f"\n--- CONTENT INTELLIGENCE (Head de Content, 30 anos exp.) ---\n{insights_buzz[:3000]}\n"
    if insights_aio:
        material += f"\n--- AI SEARCH / GEO (Especialista GEO, pioneiro) ---\n{insights_aio[:3000]}\n"

    prompt = f"""Voce e o CMO/Diretor de Estrategia responsavel por consolidar as analises de TODOS os especialistas
abaixo em um plano estrategico integrado de nivel board-level. Use frameworks McKinsey e Bain para estruturar.

{REGRAS_ANALISE}

ANALISES DOS ESPECIALISTAS:
{material}

Produza um PLANO ESTRATEGICO INTEGRADO usando os frameworks abaixo:

## 1. EXECUTIVE SUMMARY (estilo McKinsey: Situation-Complication-Resolution)
**Situacao**: Contexto de mercado e posicao atual de {prospect} (2-3 frases com numeros)
**Complicacao**: O que esta em risco e por que agir agora e urgente (2-3 frases com impacto financeiro)
**Resolucao**: A tese estrategica central -- o que propomos e o resultado esperado (2-3 frases)

## 2. DIGITAL MATURITY ASSESSMENT
Avalie {prospect} no Digital Maturity Model (escala 1-5):
| Dimensao | Nivel Atual | Nivel Alvo (12m) | Gap |
- SEO & Conteudo: [1-5]
- Social Media: [1-5]
- Performance Media: [1-5]
- Analytics & Data: [1-5]
- AI Readiness: [1-5]
- Integracao de Canais: [1-5]
**Score Geral**: X/5 -- Classificacao: Iniciante / Em Desenvolvimento / Intermediario / Avancado / Lider

## 3. HIGHLIGHTS EXECUTIVOS
Para CADA canal analisado, extraia os 3-5 insights mais impactantes.
Formato: dado concreto + "isso significa que..." + impacto estimado em R$ ou %.
Ordene por impacto financeiro (maior primeiro).

## 4. DIAGNOSTICO INTEGRADO (ANALISE CRUZADA)
A "grande narrativa" que emerge do cruzamento de TODOS os dados:
- Sinergias entre canais: onde investimento em um canal potencializa outro
- Dependencias perigosas: riscos de concentracao em canal/estrategia unica
- Correlacoes descobertas: padroes que so aparecem cruzando dados de multiplos canais
- Posicao competitiva: onde {prospect} ganha e onde perde vs mercado (mapa de calor verbal)
- Biggest Threat: qual concorrente representa maior ameaca e em quais frentes
- Biggest Opportunity: a oportunidade de maior impacto com menor competicao

## 5. MATRIZ ESTRATEGICA (ESFORCO x IMPACTO)
Classifique TODAS as recomendacoes em 4 quadrantes:

**QUICK WINS (Alto Impacto, Baixo Esforco) -- 0 a 30 dias:**
Para cada: acao especifica + KPI + meta + impacto estimado
Estas sao as prioridades #1 -- resultados rapidos que geram momentum.

**PROJETOS ESTRATEGICOS (Alto Impacto, Alto Esforco) -- 1 a 6 meses:**
Para cada: acao + recursos necessarios + KPI + meta + ROI projetado
O core do plano -- onde a maior parte do investimento deve ir.

**OTIMIZACOES INCREMENTAIS (Baixo Impacto, Baixo Esforco) -- ongoing:**
Melhorias continuas que compoem resultado ao longo do tempo.

**APOSTAS DE LONGO PRAZO (Alto Impacto, Alto Esforco) -- 6 a 12 meses:**
Iniciativas transformacionais que redefinem a presenca digital.

## 6. ROADMAP COM OKRs (12 MESES)
Estruture em OKRs por trimestre:

**Q1 -- Fundacao e Quick Wins:**
- Objective: [objetivo claro]
- KR1: [metrica mensuravel]
- KR2: [metrica mensuravel]
- KR3: [metrica mensuravel]
- Investimento estimado: R$X

**Q2 -- Escala e Otimizacao:**
[mesmo formato]

**Q3 -- Diferenciacao:**
[mesmo formato]

**Q4 -- Lideranca:**
[mesmo formato]

## 7. BUSINESS CASE SIMPLIFICADO
| Item | Valor |
- Investimento total estimado (12 meses): R$X
- Retorno projetado em trafego incremental: +X%
- Retorno projetado em leads/oportunidades: +X/mes
- Retorno projetado em receita: R$X
- ROI estimado: X:1
- Payback period: X meses

## 8. RECOMENDACOES FINAIS (TOP 7)
As 7 acoes de MAIOR impacto no negocio, em ordem de prioridade.
Para cada:
1. **O que fazer** (acao concreta e especifica)
2. **Por que** (dado que sustenta + impacto se nao agir)
3. **Como** (primeiros passos praticos)
4. **Resultado esperado** (KPI + meta + prazo)
5. **Investimento estimado** vs **Retorno projetado**"""

    return gerar_texto(prompt, especialista='estrategico')

# =============================================================================
# FUNCOES DE GERACAO DE SLIDES (DINAMICO)
# =============================================================================
def gerar_slides_completos(prospect, insights_seo, insights_social, insights_trafego, insights_midia, insights_buzz, insights_aio, recomendacoes, kw_principais, concorrentes):
    """Gera slides dinamicamente via IA com base nos insights disponíveis."""

    # Monta o inventario de insights disponiveis
    secoes_disponiveis = []
    conteudo_para_slides = ""

    if insights_seo and insights_seo.strip():
        secoes_disponiveis.append("SEO")
        conteudo_para_slides += f"\n\n=== ANALISE DE SEO (pelo Head de SEO) ===\n{insights_seo}"

    if insights_social and insights_social.strip():
        secoes_disponiveis.append("SOCIAL MEDIA")
        conteudo_para_slides += f"\n\n=== ANALISE DE SOCIAL MEDIA (pelo Head de Social Media) ===\n{insights_social}"

    if insights_trafego and insights_trafego.strip():
        secoes_disponiveis.append("FONTES DE TRAFEGO")
        conteudo_para_slides += f"\n\n=== ANALISE DE FONTES DE TRAFEGO (pelo Head de Growth) ===\n{insights_trafego}"

    if insights_midia and insights_midia.strip():
        secoes_disponiveis.append("MIDIA PAGA")
        conteudo_para_slides += f"\n\n=== ANALISE DE MIDIA PAGA (pelo Head de Performance Media) ===\n{insights_midia}"

    if insights_buzz and insights_buzz.strip():
        secoes_disponiveis.append("BUZZ MARKETING")
        conteudo_para_slides += f"\n\n=== ANALISE DE BUZZ MARKETING (pelo Head de Content Intelligence) ===\n{insights_buzz}"

    if insights_aio and insights_aio.strip():
        secoes_disponiveis.append("AIO - AI SEARCH")
        conteudo_para_slides += f"\n\n=== ANALISE DE AIO (pelo Especialista em GEO) ===\n{insights_aio}"

    if recomendacoes and recomendacoes.strip():
        conteudo_para_slides += f"\n\n=== RECOMENDACOES ESTRATEGICAS CONSOLIDADAS (pelo CMO) ===\n{recomendacoes}"

    info_contexto = f"""PROSPECT: {prospect}
CONCORRENTES: {', '.join(concorrentes)}
PALAVRAS-CHAVE: {', '.join(kw_principais)}
SECOES COM DADOS DISPONIVEIS: {', '.join(secoes_disponiveis)}"""

    prompt = f"""Voce e um diretor criativo de apresentacoes estrategicas. Sua missao e transformar as analises abaixo
em uma apresentacao de slides de diagnostico estrategico para o cliente {prospect}.

{info_contexto}

{conteudo_para_slides}

DIRETRIZES PARA A APRESENTACAO:

1. ESTRUTURA ADAPTATIVA: Crie APENAS os slides necessarios para comunicar os insights com impacto.
   Nao ha numero fixo -- pode ser 20, 40 ou 60 slides dependendo da profundidade dos dados.
   Cada slide deve ter um proposito claro. Se nao tem conteudo relevante, nao crie o slide.

2. NARRATIVA ESTRATEGICA: A apresentacao deve contar uma historia:
   - Abertura: contexto e escopo do diagnostico
   - Sumario executivo: highlights que capturam atencao imediata
   - Analises por canal: APENAS para canais com dados disponiveis ({', '.join(secoes_disponiveis)})
   - Para cada canal: scorecard > analise detalhada > insights > conclusao (problema/implicacao/solucao)
   - Visao integrada: como os canais se conectam
   - Recomendacoes priorizadas: roadmap estrategico
   - Encerramento

3. REGRAS DE CONTEUDO POR SLIDE:
   - Titulo claro e direto (maximo 10 palavras)
   - Conteudo conciso: bullet points, nao paragrafos
   - Maximo 5-7 bullets por slide. Se precisar de mais, divida em 2 slides
   - Cada insight deve ter o formato: dado/observacao + implicacao para o negocio
   - NUNCA apresente dados brutos sem interpretacao
   - Inclua indicacoes de [GRAFICO: descricao] onde fizer sentido visual

4. FORMATO DE SAIDA (siga rigorosamente):
   Cada slide deve seguir este formato exato:

   ==SLIDE==
   TITULO: [titulo do slide]
   CONTEUDO:
   [conteudo do slide, com bullets usando "-" e sub-bullets usando "  -"]
   ==FIM_SLIDE==

5. NAO INCLUA slides institucionais da Macfor (quem somos, cases, time, metodologia).
   Esses serao adicionados separadamente. Foque 100% no diagnostico e inteligencia de mercado.

Crie a apresentacao agora. Lembre-se: qualidade e profundidade dos insights importam mais que quantidade de slides."""

    resultado_ia = gerar_texto(prompt, especialista='estrategico')
    return parse_slides_ia(resultado_ia)


def parse_slides_ia(texto_ia):
    """Converte a saida da IA em lista estruturada de slides."""
    slides = []
    if not texto_ia:
        return slides

    # Divide pelos marcadores de slide
    partes = re.split(r'==\s*SLIDE\s*==', texto_ia)

    num_slide = 0
    for parte in partes:
        parte = parte.strip()
        if not parte:
            continue

        # Remove marcador de fim
        parte = re.sub(r'==\s*FIM_SLIDE\s*==', '', parte).strip()
        if not parte:
            continue

        num_slide += 1

        # Extrai titulo
        titulo = ""
        conteudo_linhas = []
        linhas = parte.split('\n')
        modo_conteudo = False

        for linha in linhas:
            linha_strip = linha.strip()
            if linha_strip.upper().startswith('TITULO:'):
                titulo = linha_strip[7:].strip()
            elif linha_strip.upper().startswith('CONTEUDO:') or linha_strip.upper().startswith('CONTEUDO :'):
                modo_conteudo = True
            elif modo_conteudo and linha_strip:
                conteudo_linhas.append(linha_strip)
            elif not titulo and not modo_conteudo and linha_strip:
                # Fallback: primeira linha nao vazia e titulo
                titulo = linha_strip
                modo_conteudo = True

        # Monta o slide
        separador = f"{'=' * 50} SLIDE {num_slide} {'=' * 50}"
        slides.append((separador,))
        if titulo:
            slides.append((titulo,))
        for linha in conteudo_linhas:
            slides.append((linha,))
        slides.append(("",))

    return slides

def gerar_documento_interno(prospect, insights_seo, insights_social, insights_trafego, insights_midia, insights_buzz, insights_aio, recomendacoes, dados_extras=""):
    """Gera documento INTERNO da agencia com todas as analises detalhadas e notas tecnicas."""

    secao_extras = ""
    if dados_extras:
        secao_extras = f"\n---\n\n## DADOS COMPLEMENTARES\n\n{dados_extras}\n"

    # Monta lista de secoes com conteudo disponivel
    secoes = []
    if insights_seo:
        secoes.append(f"## SEO\n\n{insights_seo}")
    if insights_social:
        secoes.append(f"## SOCIAL MEDIA\n\n{insights_social}")
    if insights_trafego:
        secoes.append(f"## FONTES DE TRAFEGO\n\n{insights_trafego}")
    if insights_midia:
        secoes.append(f"## MIDIA PAGA\n\n{insights_midia}")
    if insights_buzz:
        secoes.append(f"## BUZZ MARKETING / CONTENT INTELLIGENCE\n\n{insights_buzz}")
    if insights_aio:
        secoes.append(f"## AI SEARCH (GEO)\n\n{insights_aio}")

    corpo_analises = "\n\n---\n\n".join(secoes)

    documento = f"""# DOCUMENTO INTERNO -- DIAGNOSTICO {prospect.upper()}
**Data:** {datetime.now().strftime('%d/%m/%Y')}
**Uso:** INTERNO AGENCIA -- Nao compartilhar com cliente
**Elaborado por:** Equipe de Inteligencia de Mercado -- Macfor

---

{corpo_analises}
{secao_extras}
---

## RECOMENDACOES ESTRATEGICAS CONSOLIDADAS

{recomendacoes}

---

*Documento interno. Contem analises tecnicas detalhadas, notas metodologicas e dados brutos interpretados.
Para o cliente, usar o Documento de Apresentacao.*
"""
    return documento


def gerar_documento_cliente(prospect, insights_seo, insights_social, insights_trafego, insights_midia, insights_buzz, insights_aio, recomendacoes):
    """Gera documento de APRESENTACAO para o cliente via IA -- linguagem executiva, sem jargoes tecnicos."""

    # Coleta todo o material de analise
    material = ""
    if insights_seo:
        material += f"\n--- SEO ---\n{insights_seo}\n"
    if insights_social:
        material += f"\n--- SOCIAL MEDIA ---\n{insights_social}\n"
    if insights_trafego:
        material += f"\n--- TRAFEGO ---\n{insights_trafego}\n"
    if insights_midia:
        material += f"\n--- MIDIA PAGA ---\n{insights_midia}\n"
    if insights_buzz:
        material += f"\n--- BUZZ ---\n{insights_buzz}\n"
    if insights_aio:
        material += f"\n--- AI SEARCH ---\n{insights_aio}\n"
    if recomendacoes:
        material += f"\n--- RECOMENDACOES ---\n{recomendacoes}\n"

    prompt = f"""Voce e o Senior Partner da Macfor -- a agencia de marketing intelligence que {prospect} contratou
para realizar um diagnostico estrategico digital de alto nivel. Este documento sera apresentado ao CEO/CMO
de {prospect} e precisa demonstrar que a Macfor e a parceira estrategica certa para transformar a presenca
digital da marca.

MATERIAL DAS ANALISES INTERNAS (base para reescrever -- NAO copie diretamente, REINTERPRETE para o cliente):
{material}

== OBJETIVO DESTE DOCUMENTO ==
Este documento tem DOIS propositos simultaneos:
1. Entregar inteligencia de mercado de altissimo valor que o cliente NUNCA recebeu de outra agencia
2. Demonstrar a profundidade analitica e sofisticacao estrategica da Macfor como parceira

O cliente deve terminar de ler pensando: "Nunca tive acesso a esse nivel de inteligencia. Preciso dessa equipe."

== REGRAS DE CONSTRUCAO DO DOCUMENTO ==

1. LINGUAGEM DE BOARDROOM: O leitor e C-level. Nao use jargoes tecnicos sem traducao.
   NUNCA: "O Domain Authority e 35 e o crawl budget esta subotimizado"
   SEMPRE: "A credibilidade digital da marca esta 40% abaixo dos lideres do setor, o que significa
   que para cada R$1 que seus concorrentes investem em visibilidade, voces precisam investir R$1.60
   para o mesmo resultado. Isso representa um custo invisivel de R$X/mes."

2. NARRATIVA DE CONSULTOR SENIOR: Escreva como se estivesse apresentando pessoalmente ao board.
   Cada paragrafo flui naturalmente para o proximo. Nada de "Secao de Insights" separada.
   Os insights estao TECIDOS na narrativa, como descobertas que emergem organicamente da analise.

3. AUTORIDADE PELA PROFUNDIDADE: Demonstre expertise atraves de:
   - Referencias a benchmarks de mercado ("Segundo dados de 808 milhoes de dominios rastreados...")
   - Contextualizacao setorial ("No setor de [segmento], a media de mercado e X, o que posiciona voces em Y")
   - Analises cruzadas que so um especialista faria ("A queda de trafego organico combinada com o aumento
     de CPC revela uma espiral perigosa: quanto mais voces perdem organico, mais caro fica compensar com ads")
   - Projecoes fundamentadas ("Se a tendencia atual continuar, em 12 meses o concorrente A tera Z% mais
     share of voice, o que tipicamente se traduz em W pontos de market share segundo estudos de Les Binet")

4. IMPACTO FINANCEIRO EM CADA DESCOBERTA: Cada achado responde:
   - "Quanto isso esta custando hoje?" (custo de oportunidade atual)
   - "Quanto pode gerar se corrigido?" (upside potencial)
   - "O que acontece se nao agir?" (custo de inacao)

5. ESTRUTURA DO DOCUMENTO:

   # DIAGNOSTICO ESTRATEGICO DIGITAL -- {prospect}

   ## Sumario Executivo
   As 5-7 descobertas mais impactantes, cada uma em 2-3 linhas com numeros reais e impacto financeiro.
   Este sumario sozinho ja deve valer a reuniao.

   ## Posicao Competitiva Digital
   Onde {prospect} esta vs mercado. Use linguagem de "campo de batalha":
   - Onde voces lideram (celebrar e proteger)
   - Onde estao competitivos (otimizar e escalar)
   - Onde estao atras (atacar ou contornar)
   - Onde estao ausentes (oportunidades inexploradas)
   Inclua [TABELA COMPARATIVA] e [GRAFICO DE POSICAO] onde fizer sentido.

   ## Inteligencia por Canal
   APENAS canais com dados reais. Para cada canal:
   - "O que os dados revelam" (descobertas com numeros)
   - "O que isso significa para o negocio" (traducao financeira)
   - "Oportunidades identificadas" (o que pode ser capturado e quanto vale)
   Transicoes naturais entre canais, mostrando como se conectam.

   ## Visao Integrada: O Ecossistema Digital de {prospect}
   A "grande narrativa": como todos os canais se conectam, onde ha sinergias nao exploradas,
   onde ha dependencias perigosas, qual e o nivel de maturidade digital vs mercado.

   ## Plano de Acao Estrategico
   Priorizado pela Matriz Impacto x Esforco da Macfor:
   - Acoes Imediatas (0-30 dias): quick wins que geram resultados visiveis rapidamente
   - Projetos Estrategicos (1-6 meses): iniciativas de medio prazo com ROI substancial
   - Transformacao Digital (6-12 meses): projetos que redefinem a presenca digital
   Cada acao com: o que fazer, resultado esperado, KPI de sucesso.

   ## Proximo Passo
   Proposta clara do que a Macfor recomenda como proximo movimento. Tom confiante:
   "Com base neste diagnostico, recomendamos iniciar por [X] que projeta um retorno de [Y] em [Z] meses."

6. TOM INEGOCIAVEL: Consultivo, confiante, baseado em evidencias.
   NUNCA: "acreditamos", "sugerimos", "talvez", "possivelmente"
   SEMPRE: "os dados demonstram", "nossa analise revela", "recomendamos", "identificamos"
   O tom e de quem JA SABE a resposta e esta compartilhando a inteligencia.

7. EXTENSAO: Completo e profundo, mas cada paragrafo tem proposito. Se uma frase nao adiciona
   inteligencia nova, nao a inclua. Qualidade > quantidade. Cada linha deve fazer o leitor pensar
   "isso eu nao sabia" ou "isso muda minha perspectiva".

Formate em Markdown com hierarquia clara de titulos. Use **negrito** para dados-chave e descobertas criticas."""

    return gerar_texto(prompt, especialista='estrategico')

# =============================================================================
# GERACAO DE DOCX FORMATADO
# =============================================================================
COR_PRIMARIA = RGBColor(0x0D, 0x1B, 0x2A)   # dark navy
COR_SECUNDARIA = RGBColor(0x1B, 0x4D, 0x89)  # royal blue
COR_ACCENT = RGBColor(0xE8, 0x6C, 0x00)      # warm orange
COR_SUCESSO = RGBColor(0x0A, 0x8A, 0x5E)     # emerald green
COR_TEXTO = RGBColor(0x2D, 0x3A, 0x4A)       # dark slate
COR_TEXTO_LEVE = RGBColor(0x6C, 0x75, 0x7D)  # muted gray

def _setup_styles(doc):
    """Configura estilos profissionais do documento."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = COR_TEXTO
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.2

    configs = [
        (1, 22, COR_PRIMARIA, 24, 10),
        (2, 16, COR_SECUNDARIA, 18, 8),
        (3, 13, COR_PRIMARIA, 14, 6),
    ]
    for nivel, size, color, space_before, space_after in configs:
        h = doc.styles[f'Heading {nivel}']
        h.font.name = 'Calibri'
        h.font.size = Pt(size)
        h.font.color.rgb = color
        h.font.bold = True
        h.paragraph_format.space_before = Pt(space_before)
        h.paragraph_format.space_after = Pt(space_after)
        if nivel == 1:
            h.paragraph_format.keep_with_next = True

def _add_capa(doc, prospect, tipo='cliente'):
    """Adiciona capa elegante ao documento."""
    # Espacamento superior
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Linha decorativa superior
    linha = doc.add_paragraph()
    linha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = linha.add_run('_' * 40)
    run.font.color.rgb = COR_ACCENT
    run.font.size = Pt(14)

    doc.add_paragraph('')

    # Titulo principal
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo.paragraph_format.space_after = Pt(4)
    run = titulo.add_run('DIAGNOSTICO')
    run.font.size = Pt(36)
    run.font.color.rgb = COR_PRIMARIA
    run.font.bold = True
    run.font.name = 'Calibri'

    sub1 = doc.add_paragraph()
    sub1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub1.paragraph_format.space_after = Pt(20)
    run = sub1.add_run('ESTRATEGICO DIGITAL')
    run.font.size = Pt(28)
    run.font.color.rgb = COR_SECUNDARIA
    run.font.bold = False
    run.font.name = 'Calibri'

    # Nome do prospect
    nome_p = doc.add_paragraph()
    nome_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nome_p.paragraph_format.space_after = Pt(30)
    run = nome_p.add_run(prospect.upper())
    run.font.size = Pt(24)
    run.font.color.rgb = COR_ACCENT
    run.font.bold = True
    run.font.name = 'Calibri'

    # Linha decorativa inferior
    linha2 = doc.add_paragraph()
    linha2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = linha2.add_run('_' * 40)
    run.font.color.rgb = COR_ACCENT
    run.font.size = Pt(14)

    doc.add_paragraph('')
    doc.add_paragraph('')

    # Tipo de documento
    label = 'Documento Interno | Uso Exclusivo da Agencia' if tipo == 'interno' else 'Apresentacao Executiva | Confidencial'
    tipo_p = doc.add_paragraph()
    tipo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tipo_p.add_run(label)
    run.font.size = Pt(11)
    run.font.color.rgb = COR_TEXTO_LEVE
    run.font.name = 'Calibri'
    run.italic = True

    # Data
    data_p = doc.add_paragraph()
    data_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = data_p.add_run(datetime.now().strftime('%B %Y').title())
    run.font.size = Pt(12)
    run.font.color.rgb = COR_TEXTO_LEVE
    run.font.name = 'Calibri'

    doc.add_paragraph('')
    doc.add_paragraph('')

    # Marca Macfor
    marca = doc.add_paragraph()
    marca.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = marca.add_run('MACFOR')
    run.font.size = Pt(16)
    run.font.color.rgb = COR_PRIMARIA
    run.font.bold = True
    run.font.name = 'Calibri'

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run('Marketing Intelligence & Digital Strategy')
    run.font.size = Pt(10)
    run.font.color.rgb = COR_TEXTO_LEVE
    run.font.name = 'Calibri'
    run.italic = True

    doc.add_page_break()

def _colorir_celula(cell, cor_hex):
    """Aplica cor de fundo a uma celula."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{cor_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def _formatar_celula(cell, font_size=9, bold=False, color=None, align=None):
    """Formata texto de uma celula."""
    for p in cell.paragraphs:
        if align:
            p.alignment = align
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            r.font.size = Pt(font_size)
            r.font.name = 'Calibri'
            r.font.bold = bold
            if color:
                r.font.color.rgb = color

def _add_tabela_markdown(doc, linhas_tabela):
    """Converte tabela markdown em tabela DOCX elegante."""
    # Parse das linhas da tabela
    rows_data = []
    for linha in linhas_tabela:
        cells = [c.strip() for c in linha.strip('|').split('|')]
        # Ignora linhas separadoras (---|---|---)
        if cells and not all(re.match(r'^[-:]+$', c) for c in cells):
            rows_data.append(cells)

    if len(rows_data) < 1:
        return

    num_cols = max(len(r) for r in rows_data)
    table = doc.add_table(rows=0, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows_data):
        row = table.add_row()
        for j in range(num_cols):
            val = row_data[j] if j < len(row_data) else ''
            row.cells[j].text = val
            if i == 0:
                # Header row
                _colorir_celula(row.cells[j], '0D1B2A')
                _formatar_celula(row.cells[j], font_size=9, bold=True,
                                color=RGBColor(0xFF, 0xFF, 0xFF))
            else:
                if i % 2 == 0:
                    _colorir_celula(row.cells[j], 'F4F6F9')
                _formatar_celula(row.cells[j], font_size=9)

    doc.add_paragraph('')

def _add_runs_formatados(paragraph, texto):
    """Adiciona texto com **bold** e *italic* como runs formatados."""
    partes = re.split(r'(\*\*.*?\*\*|\*.*?\*)', texto)
    for parte in partes:
        if parte.startswith('**') and parte.endswith('**'):
            run = paragraph.add_run(parte[2:-2])
            run.bold = True
            run.font.color.rgb = COR_PRIMARIA
        elif parte.startswith('*') and parte.endswith('*'):
            run = paragraph.add_run(parte[1:-1])
            run.italic = True
            run.font.color.rgb = COR_TEXTO_LEVE
        else:
            paragraph.add_run(parte)

def _add_callout_box(doc, texto, tipo='info'):
    """Adiciona caixa de destaque visual."""
    cores = {
        'info': ('E8F0FE', COR_SECUNDARIA),
        'warning': ('FFF3E0', COR_ACCENT),
        'success': ('E8F5E9', COR_SUCESSO),
    }
    bg, _ = cores.get(tipo, cores['info'])

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    _colorir_celula(cell, bg)
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    _add_runs_formatados(p, texto)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'

    doc.add_paragraph('')

def _markdown_para_docx(doc, texto_md):
    """Converte markdown em DOCX com suporte a tabelas, callouts e formatacao rica."""
    if not texto_md:
        return

    linhas = texto_md.split('\n')
    i = 0
    tabela_buffer = []
    em_tabela = False

    while i < len(linhas):
        linha = linhas[i]
        stripped = linha.strip()

        # Detecta tabela markdown
        if '|' in stripped and stripped.startswith('|'):
            tabela_buffer.append(stripped)
            em_tabela = True
            i += 1
            continue
        elif em_tabela:
            # Acabou a tabela
            _add_tabela_markdown(doc, tabela_buffer)
            tabela_buffer = []
            em_tabela = False

        if not stripped:
            i += 1
            continue

        # Headings
        if stripped.startswith('#### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(stripped[5:])
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = COR_SECUNDARIA
            run.font.name = 'Calibri'
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)
        # Separadores
        elif stripped.startswith('---'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
        # Blockquotes como callout
        elif stripped.startswith('> '):
            texto_quote = stripped[2:]
            _add_callout_box(doc, texto_quote, tipo='info')
        # Bullet list
        elif stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            _add_runs_formatados(p, stripped[2:])
        # Sub-bullet
        elif re.match(r'^  +[-*] ', stripped):
            p = doc.add_paragraph(style='List Bullet 2')
            texto = re.sub(r'^  +[-*] ', '', stripped)
            _add_runs_formatados(p, texto)
        # Numbered list
        elif re.match(r'^\d+[\.\)] ', stripped):
            p = doc.add_paragraph(style='List Number')
            texto = re.sub(r'^\d+[\.\)] ', '', stripped)
            _add_runs_formatados(p, texto)
        # Normal paragraph
        else:
            p = doc.add_paragraph()
            _add_runs_formatados(p, stripped)

        i += 1

    # Flush tabela restante
    if tabela_buffer:
        _add_tabela_markdown(doc, tabela_buffer)

def _add_rodape(doc):
    """Adiciona rodape elegante."""
    doc.add_paragraph('')
    # Linha separadora
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sep.add_run('_' * 50)
    run.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)
    run.font.size = Pt(8)

    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('MACFOR | Marketing Intelligence & Digital Strategy')
    run.font.size = Pt(9)
    run.font.color.rgb = COR_TEXTO_LEVE
    run.font.name = 'Calibri'
    run.bold = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run('Este documento e confidencial e destinado exclusivamente ao uso do destinatario.')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    run.font.name = 'Calibri'
    run.italic = True

def gerar_docx(prospect, texto_markdown, dados_brutos=None, tipo='cliente'):
    """Gera documento DOCX profissional."""
    _ = dados_brutos  # Mantido na assinatura por compatibilidade
    doc = Document()

    # Margens elegantes
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    _setup_styles(doc)
    _add_capa(doc, prospect, tipo)

    # Conteudo principal
    _markdown_para_docx(doc, texto_markdown)

    # Rodape
    _add_rodape(doc)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# =============================================================================
# INTERFACE PRINCIPAL
# =============================================================================
st.sidebar.header("Configuração")

# Informacoes basicas
st.sidebar.subheader("Informacoes do Cliente")
input_prospect = st.sidebar.text_input("Nome do Prospect/Cliente", placeholder="Ex: Empresa XYZ", key="input_prospect")
input_concorrentes = st.sidebar.text_input("Concorrentes (separados por virgula)", placeholder="Ex: Concorrente A, Concorrente B", key="input_concorrentes")
input_kw = st.sidebar.text_input("Palavras-chave principais (separadas por virgula)", placeholder="Ex: colchao, colchao ortopedico, cama box", key="input_kw")

st.sidebar.markdown("---")

# Upload dos CSVs
st.sidebar.subheader("Upload de Arquivos")
st.sidebar.caption("Cada campo aceita multiplos arquivos. Nenhum campo e obrigatorio -- carregue apenas o que tiver.")
arquivos_seo = st.sidebar.file_uploader("SEO (historico, rank, trafego organico)", type=['csv'], accept_multiple_files=True, key="seo")
arquivos_keywords = st.sidebar.file_uploader("Keywords (ranking, volume, analise)", type=['csv'], accept_multiple_files=True, key="kw")
arquivos_facebook = st.sidebar.file_uploader("Facebook", type=['csv'], accept_multiple_files=True, key="fb")
arquivos_instagram = st.sidebar.file_uploader("Instagram", type=['csv'], accept_multiple_files=True, key="ig")
arquivos_tiktok = st.sidebar.file_uploader("TikTok", type=['csv'], accept_multiple_files=True, key="tt")
arquivos_linkedin = st.sidebar.file_uploader("LinkedIn", type=['csv'], accept_multiple_files=True, key="li")
arquivos_whatsapp = st.sidebar.file_uploader("WhatsApp", type=['csv'], accept_multiple_files=True, key="wa")
arquivos_concorrencia = st.sidebar.file_uploader("Concorrencia (benchmarks, comparativos)", type=['csv'], accept_multiple_files=True, key="conc")
arquivos_autoridade = st.sidebar.file_uploader("Autoridade / Dominio (backlinks, DA)", type=['csv'], accept_multiple_files=True, key="auth")

st.sidebar.markdown("---")
st.sidebar.subheader("Dados Adicionais")
st.sidebar.caption("Qualquer CSV extra -- a IA interpretara automaticamente.")
arquivos_extras = st.sidebar.file_uploader("Adicionais", type=['csv'], accept_multiple_files=True, key="extras")

st.sidebar.markdown("---")
st.sidebar.subheader("Contexto Adicional")
contexto_usuario = st.sidebar.text_area(
    "Informacoes extras para a analise",
    placeholder="Ex: O cliente atua no segmento B2B de colchoes premium, faturamento anual de R$50M, objetivo principal e aumentar market share no digital frente ao concorrente X que lancou e-commerce recentemente...",
    height=150,
    key="contexto_usuario"
)

# Botão para processar
if st.sidebar.button("Processar Dados e Gerar Diagnóstico"):
    # Salvar contexto adicional no session_state para uso nos prompts
    st.session_state.contexto_adicional = contexto_usuario.strip() if contexto_usuario else ""

    with st.spinner("Carregando e processando dados..."):
        dados_brutos = {}

        # Carregar todos os CSVs e armazenar brutos
        mapa_uploads = {
            'seo': arquivos_seo,
            'keywords': arquivos_keywords,
            'facebook': arquivos_facebook,
            'instagram': arquivos_instagram,
            'tiktok': arquivos_tiktok,
            'linkedin': arquivos_linkedin,
            'whatsapp': arquivos_whatsapp,
            'concorrencia': arquivos_concorrencia,
            'autoridade': arquivos_autoridade,
        }
        for chave, arquivos in mapa_uploads.items():
            if arquivos:
                df = carregar_e_combinar_csvs(arquivos)
                if df is not None and not df.empty:
                    dados_brutos[chave] = df

        # Prospect e concorrentes: prioriza input manual, fallback para CSV
        prospect_manual = input_prospect.strip() if input_prospect else ""
        conc_manual = [c.strip() for c in input_concorrentes.split(',') if c.strip()] if input_concorrentes else []
        kw_manual = [k.strip() for k in input_kw.split(',') if k.strip()] if input_kw else []

        if prospect_manual:
            st.session_state.nome_prospect = prospect_manual
        elif 'concorrencia' in dados_brutos:
            prospect, _ = processar_dados_prospect_concorrentes(dados_brutos['concorrencia'])
            st.session_state.nome_prospect = prospect
        else:
            st.session_state.nome_prospect = "Prospect"

        if conc_manual:
            st.session_state.concorrentes = conc_manual
        elif 'concorrencia' in dados_brutos:
            _, concorrentes = processar_dados_prospect_concorrentes(dados_brutos['concorrencia'])
            st.session_state.concorrentes = concorrentes
        else:
            st.session_state.concorrentes = []

        if kw_manual:
            st.session_state.kw_principais = kw_manual
        elif 'concorrencia' in dados_brutos:
            kws = processar_kw_principais(dados_brutos['concorrencia'])
            st.session_state.kw_principais = kws if kws else []
        else:
            st.session_state.kw_principais = []

        # CSVs extras
        if arquivos_extras:
            for i, arq in enumerate(arquivos_extras):
                df_extra = carregar_csv(arq)
                if df_extra is not None and not df_extra.empty:
                    nome = arq.name.replace('.csv', '').replace('.CSV', '')
                    dados_brutos[f'extra_{nome}'] = df_extra

        st.session_state.dados_brutos = dados_brutos
        st.session_state.dados_carregados = True
        st.success(f"Dados carregados: {len(dados_brutos)} conjuntos de dados ({', '.join(dados_brutos.keys())})")
    
    with st.spinner("Gerando insights com IA -- isso pode levar alguns minutos..."):
        prospect = st.session_state.nome_prospect
        concorrentes = st.session_state.concorrentes
        kw_principais = st.session_state.kw_principais
        db = st.session_state.dados_brutos

        # Monta contexto de dados brutos por area
        dados_seo = ""
        for k in ['seo', 'keywords', 'autoridade', 'concorrencia']:
            if k in db:
                dados_seo += f"\n--- CSV: {k} ---\n{df_para_contexto(db[k])}\n"

        dados_social = ""
        for k in ['facebook', 'instagram', 'tiktok', 'linkedin', 'whatsapp']:
            if k in db:
                dados_social += f"\n--- CSV: {k} ---\n{df_para_contexto(db[k])}\n"

        dados_todos = ""
        for k, v in db.items():
            dados_todos += f"\n--- CSV: {k} ---\n{df_para_contexto(v, 25)}\n"

        dados_extras = ""
        for k in db:
            if k.startswith('extra_'):
                dados_extras += f"\n--- {k} ---\n{df_para_contexto(db[k])}\n"

        # SEO -- gera se houver dados de SEO/keywords/autoridade OU contexto do usuario
        if dados_seo or st.session_state.contexto_adicional:
            st.info("Analisando SEO...")
            st.session_state.insights_seo = gerar_insights_seo(
                dados_seo or "(Nenhum CSV de SEO fornecido. Use o contexto adicional e dados disponiveis.)",
                kw_principais, prospect, concorrentes
            )

        # Social Media -- gera se houver dados sociais
        if dados_social:
            st.info("Analisando Social Media...")
            st.session_state.insights_social = gerar_insights_social(
                dados_social, prospect, concorrentes
            )

        # Tráfego -- gera se houver qualquer dado
        if dados_todos:
            st.info("Analisando Fontes de Trafego...")
            st.session_state.insights_trafego = gerar_insights_trafego(
                dados_todos, prospect, concorrentes
            )

        # Mídia Paga -- gera se houver qualquer dado
        if dados_todos:
            st.info("Analisando Midia Paga...")
            st.session_state.insights_midia_paga = gerar_insights_midia_paga(
                dados_todos, prospect, concorrentes
            )

        # Buzz -- sempre gera (usa conhecimento de mercado + contexto)
        st.info("Analisando Buzz Marketing...")
        st.session_state.insights_buzz = gerar_insights_buzz(
            prospect, kw_principais
        )

        # AIO -- gera se houver dados de SEO ou keywords
        if dados_seo or kw_principais:
            st.info("Analisando AI Search / AIO...")
            st.session_state.insights_aio = gerar_insights_aio(
                dados_seo or "(Sem dados de SEO. Analise com base nas keywords e contexto.)",
                prospect, kw_principais
            )

        # Recomendações consolidadas -- sempre gera com o que tiver
        st.info("Consolidando recomendacoes estrategicas...")
        st.session_state.recomendacoes = gerar_recomendacoes_estrategicas(
            st.session_state.insights_seo or "",
            st.session_state.insights_social or "",
            st.session_state.insights_trafego or "",
            st.session_state.insights_midia_paga or "",
            st.session_state.insights_buzz or "",
            st.session_state.insights_aio or "",
            prospect
        )

        # Gerar slides
        st.info("Gerando apresentacao de slides...")
        st.session_state.slides_gerados = gerar_slides_completos(
            prospect,
            st.session_state.insights_seo or "",
            st.session_state.insights_social or "",
            st.session_state.insights_trafego or "",
            st.session_state.insights_midia_paga or "",
            st.session_state.insights_buzz or "",
            st.session_state.insights_aio or "",
            st.session_state.recomendacoes or "",
            kw_principais,
            concorrentes
        )

        # Gerar documento INTERNO (agencia)
        st.info("Gerando documento interno...")
        st.session_state.documento_analise = gerar_documento_interno(
            prospect,
            st.session_state.insights_seo or "",
            st.session_state.insights_social or "",
            st.session_state.insights_trafego or "",
            st.session_state.insights_midia_paga or "",
            st.session_state.insights_buzz or "",
            st.session_state.insights_aio or "",
            st.session_state.recomendacoes or "",
            dados_extras
        )

        # Gerar documento do CLIENTE (apresentacao executiva)
        st.info("Gerando documento de apresentacao para o cliente...")
        st.session_state.documento_cliente = gerar_documento_cliente(
            prospect,
            st.session_state.insights_seo or "",
            st.session_state.insights_social or "",
            st.session_state.insights_trafego or "",
            st.session_state.insights_midia_paga or "",
            st.session_state.insights_buzz or "",
            st.session_state.insights_aio or "",
            st.session_state.recomendacoes or ""
        )

        st.session_state.relatorio_gerado = True
        st.success("Diagnostico gerado com sucesso!")

# =============================================================================
# EXIBIÇÃO DOS RESULTADOS
# =============================================================================
if st.session_state.relatorio_gerado:
    st.markdown("---")
    
    nome_safe = st.session_state.nome_prospect.lower().replace(' ', '_')
    data_str = datetime.now().strftime('%Y%m%d')

    tab_cliente, tab_interno, tab_slides, tab_dados = st.tabs([
        "Apresentacao Cliente",
        "Documento Interno",
        "Slides",
        "Dados Brutos"
    ])

    with tab_cliente:
        st.header(f"Diagnostico Estrategico -- {st.session_state.nome_prospect}")
        st.caption("Documento para apresentar ao cliente. Linguagem executiva, foco em negocio.")
        st.markdown(st.session_state.documento_cliente)
        col_dl1, col_dl2 = st.columns(2)
        with col_dl1:
            st.download_button(
                label="Baixar Markdown",
                data=st.session_state.documento_cliente,
                file_name=f"cliente_{nome_safe}_{data_str}.md",
                mime="text/markdown",
                key="dl_cliente_md"
            )
        with col_dl2:
            docx_cliente = gerar_docx(
                st.session_state.nome_prospect,
                st.session_state.documento_cliente,
                st.session_state.dados_brutos,
                tipo='cliente'
            )
            st.download_button(
                label="Baixar DOCX Formatado",
                data=docx_cliente,
                file_name=f"cliente_{nome_safe}_{data_str}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_cliente_docx"
            )

    with tab_interno:
        st.header(f"Documento Interno -- {st.session_state.nome_prospect}")
        st.caption("USO INTERNO DA AGENCIA. Analises detalhadas, dados brutos, notas tecnicas.")
        st.markdown(st.session_state.documento_analise)
        col_dl3, col_dl4 = st.columns(2)
        with col_dl3:
            st.download_button(
                label="Baixar Markdown",
                data=st.session_state.documento_analise,
                file_name=f"interno_{nome_safe}_{data_str}.md",
                mime="text/markdown",
                key="dl_interno_md"
            )
        with col_dl4:
            docx_interno = gerar_docx(
                st.session_state.nome_prospect,
                st.session_state.documento_analise,
                st.session_state.dados_brutos,
                tipo='interno'
            )
            st.download_button(
                label="Baixar DOCX Formatado",
                data=docx_interno,
                file_name=f"interno_{nome_safe}_{data_str}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_interno_docx"
            )
    
    with tab_slides:
        st.header("Slides do Diagnostico")

        slides_list = st.session_state.slides_gerados
        if slides_list:
            # Conta slides reais (pelos separadores)
            total_slides = sum(1 for s in slides_list if s and s[0] and 'SLIDE' in s[0] and '===' in s[0])
            st.markdown(f"**{total_slides} slides gerados** com base nos dados disponiveis.")
            st.markdown("---")

            # Renderiza cada slide como card
            slide_atual = None
            conteudo_slide = []
            for elemento in slides_list:
                linha = elemento[0] if isinstance(elemento, tuple) and elemento else str(elemento)
                if 'SLIDE' in linha and '===' in linha:
                    # Renderiza slide anterior
                    if slide_atual is not None:
                        with st.expander(slide_atual, expanded=False):
                            st.markdown('\n'.join(conteudo_slide))
                    slide_atual = linha.strip('= ').strip()
                    conteudo_slide = []
                elif linha.strip():
                    conteudo_slide.append(linha)

            # Renderiza ultimo slide
            if slide_atual is not None and conteudo_slide:
                with st.expander(slide_atual, expanded=False):
                    st.markdown('\n'.join(conteudo_slide))

        # Texto para download
        texto_slides = ""
        for elemento in slides_list:
            if isinstance(elemento, tuple):
                texto_slides += "\n".join(elemento) + "\n"
            else:
                texto_slides += str(elemento) + "\n"

        st.download_button(
            label="Baixar Slides (TXT)",
            data=texto_slides,
            file_name=f"slides_{st.session_state.nome_prospect.lower().replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.txt",
            mime="text/plain"
        )
    
    with tab_dados:
        st.header("Dados Processados")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Prospect e Concorrentes")
            st.write(f"**Prospect:** {st.session_state.nome_prospect}")
            st.write("**Concorrentes:**")
            for i, conc in enumerate(st.session_state.concorrentes, 1):
                st.write(f"{i}. {conc}")
            
            st.subheader("Palavras-chave Principais")
            for i, kw in enumerate(st.session_state.kw_principais, 1):
                st.write(f"{i}. {kw}")
        
        with col2:
            st.subheader("Conjuntos de Dados Carregados")
            for nome, df in st.session_state.dados_brutos.items():
                st.write(f"**{nome}**: {df.shape[0]} linhas x {df.shape[1]} colunas")

        # Mostra amostra de cada dataset
        for nome, df in st.session_state.dados_brutos.items():
            with st.expander(f"Dados: {nome} ({df.shape[0]}x{df.shape[1]})"):
                st.dataframe(df.head(15))

else:
    # Tela inicial
    st.info("👈 Carregue os arquivos CSV na barra lateral para gerar o diagnóstico.")
    
    st.markdown("""
    ## Instrucoes

    1. **Carregue os arquivos CSV** na barra lateral (nao e necessario preencher todos os campos)
    2. **Preencha o contexto adicional** com informacoes sobre o cliente e objetivos
    3. **Clique em "Processar Dados e Gerar Diagnostico"**
    4. **Baixe os resultados** em Markdown ou DOCX formatado

    ### Campos de Upload

    | Campo | O que carregar |
    |-------|---------------|
    | **SEO** | Historico SEO, trafego organico, rank tracking (SEMrush, Ahrefs) |
    | **Keywords** | Ranking de keywords, volumes, CPC, analise de oportunidades |
    | **Facebook** | Metricas de pagina, engajamento, posts |
    | **Instagram** | Followers, engagement rate, posts, stories |
    | **TikTok** | Views, engajamento, seguidores, videos |
    | **LinkedIn** | Company page, followers, posts, engajamento |
    | **WhatsApp** | Metricas de atendimento, conversao, volume |
    | **Concorrencia** | Benchmarks, dados comparativos entre players |
    | **Autoridade/Dominio** | Domain Authority, backlinks, dominios referenciadores |
    | **Adicionais** | Qualquer CSV extra (CRM, GA, vendas, NPS, etc.) |

    ### Saidas Geradas

    - **Apresentacao Cliente**: documento executivo (DOCX formatado com graficos e tabelas)
    - **Documento Interno**: analise tecnica detalhada para uso da agencia
    - **Slides**: apresentacao dinamica gerada com base nos dados
    - **Dados Brutos**: visualizacao dos datasets carregados
    """)

# Rodapé
st.markdown("---")
st.markdown("**Gerador de Diagnóstico Estratégico** | Desenvolvido para Macfor")
